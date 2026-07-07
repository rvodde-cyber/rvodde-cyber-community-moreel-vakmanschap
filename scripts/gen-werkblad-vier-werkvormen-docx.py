"""Genereer MV_20 — Vier werkvormen met gesprekskaarten (NL + EN)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "public/downloads/gesprekskaarten"

C_TITLE = RGBColor(0x1F, 0x2D, 0x54)
C_THEMES = RGBColor(0x8B, 0x69, 0x14)
C_BODY = RGBColor(0x4A, 0x55, 0x68)
C_MUTED = RGBColor(0x5F, 0x5E, 0x5A)
PHASE_COLORS = {
    "zien": RGBColor(0x18, 0x5F, 0xA5),
    "voelen": RGBColor(0x85, 0x4F, 0x0B),
    "wegen": RGBColor(0x99, 0x35, 0x56),
    "handelen": RGBColor(0x0F, 0x6E, 0x56),
}

CONTENT = {
    "nl": {
        "filename": "MV_20_VierWerkvormen_NL.docx",
        "header_mid": "MV_20  —  Vier Werkvormen met Gesprekskaarten\nReeks Moreel Vakmanschap",
        "header_right": "Zien · Voelen · Wegen · Handelen",
        "title": "Vier Werkvormen met Gesprekskaarten",
        "subtitle": "Gesprekskaarten  ·  Zien · Voelen · Wegen · Handelen  ·  Moreel vakmanschap in de praktijk",
        "intro_title": "Waarom deze indeling?",
        "intro_p1": (
            "Elke fase van moreel handelen vraagt om een ander soort oefening. Zien vraagt om scherper waarnemen "
            "vóórdat je oordeelt; Voelen vraagt om een breder gevoelsrepertoire dan je eigen automatische reactie; "
            "Wegen vraagt om zorgvuldig verschillende argumenten tegen elkaar afzetten; en Handelen vraagt om het "
            "concreet maken van wat je zou doen of zeggen. Onderstaande vier werkvormen gebruiken elk hun eigen "
            "kaartenset, precies afgestemd op wat die specifieke fase vraagt."
        ),
        "intro_p2": (
            "Deze werkvormen zijn geschikt voor groepen van circa 4 tot 8 deelnemers en duren elk 15 tot 25 minuten. "
            "Ze kunnen los worden ingezet, of achter elkaar als een volledige doorloop van een dilemma — van waarneming "
            "tot actie. Combineer ze met gesprekskaarten uit de bibliotheek en, waar passend, met werkvormen over Wegen "
            "(Moreel Beraad, Goede redenen) en Volhouden."
        ),
        "phases": [
            {
                "key": "zien",
                "label": "FASE: ZIEN",
                "name": "De Kaartenwaaier",
                "doel": "Doel: Scherpt de morele waarneming — wat merk je eigenlijk op vóórdat je gaat oordelen, en wat mis je meestal?",
                "cards": (
                    "Benodigde kaarten: een set signaalkaarten met korte waarnemingsvragen, bijvoorbeeld: "
                    "'Wie is hier stil gebleven?', 'Welk detail zou je normaal gesproken negeren?', "
                    "'Wat zou een buitenstaander direct opvallen dat jij over het hoofd ziet?', "
                    "'Welke aanname doe je hier zonder het te merken?'"
                ),
                "steps": [
                    "De begeleider leest een bewust onvolledige of ambigue casusbeschrijving voor — met opzet zonder duidelijke 'hoofdpersoon' of moraal.",
                    "Om de beurt trekt een deelnemer blind een signaalkaart uit de waaier en benoemt, uitsluitend gebaseerd op wat er letterlijk in de casus staat, een waarneming die nog niet door iemand anders is genoemd.",
                    "Herhaal tot de waaier op is of tot er geen nieuwe waarnemingen meer opduiken.",
                    "Bespreek na afloop: welke waarneming kwam het laatst, en waarom duurde het zo lang voordat iemand die opmerkte?",
                ],
                "tip": "Praktijktip — Sta niet toe dat deelnemers meteen een oordeel of oplossing koppelen aan hun waarneming. Bij deze werkvorm mag er alleen worden benoemd wat je ziet, nog niet wat je ervan vindt.",
            },
            {
                "key": "voelen",
                "label": "FASE: VOELEN",
                "name": "De Empathiekaarten",
                "doel": "Doel: Verbreedt het gevoelsrepertoire — welke emoties zou déze situatie kunnen oproepen, ook als het niet je eigen eerste reactie is?",
                "cards": (
                    "Benodigde kaarten: een set emotiekaarten met telkens één enkel gevoel — zonder uitleg — "
                    "zoals: onmacht, opluchting, schaamte, irritatie, twijfel, trots, angst, verdriet."
                ),
                "steps": [
                    "Na het horen van de casus trekt iedere deelnemer blind één emotiekaart.",
                    "Elke deelnemer legt uit waar of bij wie in de casus die specifieke emotie zou kunnen opduiken — ook als het niet de emotie is die zij zelf als eerste voelden.",
                    "De groep mag aanvullen: bij wie in de casus zou dit gevoel nóg sterker kunnen spelen?",
                    "Sluit af met een korte ronde: welke emotie op tafel had jij zelf nog niet genoemd, maar herken je nu wel?",
                ],
                "tip": "Praktijktip — Laat nadrukkelijk ook 'onprettige' of ongepast klinkende emoties toe (opluchting dat het een ander overkwam, bijvoorbeeld). Het doel is eerlijke herkenning, niet een sociaal wenselijk gevoel.",
            },
            {
                "key": "wegen",
                "label": "FASE: WEGEN",
                "name": "De Weegschaal-vloer",
                "doel": "Doel: Maakt het afwegen van argumenten fysiek en zichtbaar, in plaats van dat het alleen abstract in iemands hoofd blijft.",
                "cards": (
                    "Benodigde kaarten: losse argumentkaarten (voorbedrukt of ter plekke ingevuld) met korte argumenten "
                    "vóór en tegen een mogelijke keuze in de casus."
                ),
                "steps": [
                    "Markeer een denkbeeldige lijn op de vloer of een lange tafel, van 'volledig mee eens' aan het ene uiterste tot 'volledig mee oneens' aan het andere.",
                    "Elke deelnemer krijgt één of meerdere argumentkaarten en plaatst deze fysiek op de lijn, op basis van hoe zwaar dat argument voor hen weegt.",
                    "Zodra alle kaarten liggen, mogen deelnemers elkaars plaatsing bevragen: 'Waarom ligt jouw kaart dichter naar het midden dan die van mij?'",
                    "Sluit af met de vraag: als je de kaarten nu moest samenvoegen tot één gezamenlijk oordeel, waar zou het geheel dan ongeveer liggen?",
                ],
                "tip": "Praktijktip — Sta bewust toe dat kaarten dicht bij het midden komen te liggen. Een dilemma dat zich makkelijk aan één kant van de lijn laat oplossen, is meestal geen goed dilemma.",
            },
            {
                "key": "handelen",
                "label": "FASE: HANDELEN",
                "name": "De Actiezinnenset",
                "doel": "Doel: Overbrugt de kloof tussen weten wat goed zou zijn en het daadwerkelijk concreet zeggen of doen — vaak de zwakste schakel in moreel gedrag.",
                "cards": (
                    "Benodigde kaarten: kaarten met de aanzet van een actiezin, bijvoorbeeld: "
                    "'Het eerste wat ik nu zou doen is…', 'Wat ik hardop tegen [naam] zou zeggen is…', "
                    "'Het eerste telefoontje dat ik zou plegen is naar…', 'Als ik hier nu voor zou staan, zou ik letterlijk zeggen…'"
                ),
                "steps": [
                    "Na het bespreken van het dilemma trekt elke deelnemer blind een actiezin-kaart.",
                    "De deelnemer moet de zin hardop afmaken — concreet en specifiek (een naam, een moment, een letterlijke formulering), niet in algemeenheden als 'ik zou het bespreekbaar maken'.",
                    "De groep mag doorvragen zodra het antwoord te vaag blijft: 'Wat zeg je dan precies, en tegen wie?'",
                    "Bespreek tot slot: welke actiezin was het makkelijkst om te bedenken, en welke het moeilijkst om ook echt hardop uit te spreken?",
                ],
                "tip": "Praktijktip — Wees streng op concreetheid: 'ik zou er iets van zeggen' is geen geldig antwoord. Vraag door tot de deelnemer een zin uitspreekt die hij of zij morgen letterlijk zou kunnen gebruiken.",
            },
        ],
        "overview_title": "Overzicht in één oogopslag",
        "overview_headers": ["Fase", "Werkvorm", "Kaarten nodig"],
        "overview_rows": [
            ("ZIEN", "De Kaartenwaaier", "Signaalkaarten met korte waarnemingsvragen"),
            ("VOELEN", "De Empathiekaarten", "Emotiekaarten — één gevoel per kaart, zonder uitleg"),
            ("WEGEN", "De Weegschaal-vloer", "Argumentkaarten vóór en tegen een mogelijke keuze"),
            ("HANDELEN", "De Actiezinnenset", "Kaarten met de aanzet van een actiezin"),
        ],
        "teacher_title": "Docentennotitie — veilige leeromgeving",
        "teacher_text": (
            "Laat deelnemers bij Zien en Voelen nog geen oplossing of oordeel formuleren — dat komt pas bij Wegen en Handelen. "
            "Het door elkaar laten lopen van de fasen is de meest voorkomende valkuil: een groep die te snel naar 'wat moeten we doen' springt, "
            "slaagt vaak de waarneming en het invoelen over, en mist daardoor juist de signalen die het dilemma het lastigst maken."
        ),
        "tip_combo": "Tip: Dit werkblad werkt goed in combinatie met gesprekskaarten uit de bibliotheek — gebruik een casuskaart als gemeenschappelijke casus voor alle vier werkvormen.",
        "sources": (
            "Bronnen: Rest, J.R. (1986), Moral Development: Advances in Research and Theory  ·  "
            "Kidder, R.M. (2005), How Good People Make Tough Choices  ·  "
            "Tjin A Djie & Zwaan (2010), Managen van diversiteit op de werkvloer"
        ),
        "footer": "Fontys Hogescholen — Lectoraat Ethisch Werken · Community Moreel Vakmanschap",
        "how_label": "Hoe het werkt",
    },
    "en": {
        "filename": "MV_20_FourWorkForms_EN.docx",
        "header_mid": "MV_20  —  Four Work Forms with Conversation Cards\nMoral Craftsmanship Series",
        "header_right": "Seeing · Feeling · Weighing · Acting",
        "title": "Four Work Forms with Conversation Cards",
        "subtitle": "Conversation cards  ·  Seeing · Feeling · Weighing · Acting  ·  Moral craftsmanship in practice",
        "intro_title": "Why this structure?",
        "intro_p1": (
            "Each phase of moral action calls for a different kind of exercise. Seeing asks for sharper observation before you judge; "
            "Feeling asks for a broader emotional repertoire than your own automatic reaction; Weighing asks for carefully comparing "
            "different arguments; and Acting asks you to make concrete what you would do or say. The four work forms below each use "
            "their own card set, precisely aligned with what that specific phase requires."
        ),
        "intro_p2": (
            "These work forms suit groups of roughly 4 to 8 participants and each take 15 to 25 minutes. They can be used separately, "
            "or in sequence as a full run through a dilemma — from observation to action. Combine them with conversation cards from "
            "the library and, where appropriate, with work forms on Weighing (Moral Deliberation, Good Reasons) and Persisting."
        ),
        "phases": [
            {
                "key": "zien",
                "label": "PHASE: SEEING",
                "name": "The Card Fan",
                "doel": "Aim: Sharpens moral observation — what do you actually notice before you judge, and what do you usually miss?",
                "cards": (
                    "Cards needed: a set of signal cards with short observation prompts, for example: "
                    "'Who has remained silent here?', 'Which detail would you normally ignore?', "
                    "'What would an outsider immediately notice that you overlook?', "
                    "'What assumption are you making here without noticing?'"
                ),
                "steps": [
                    "The facilitator reads an deliberately incomplete or ambiguous case description — intentionally without a clear 'main character' or moral.",
                    "In turn, a participant draws a signal card blindly from the fan and names, based solely on what is literally in the case, an observation not yet mentioned by anyone else.",
                    "Repeat until the fan is empty or no new observations emerge.",
                    "Discuss afterwards: which observation came last, and why did it take so long before someone mentioned it?",
                ],
                "tip": "Practical tip — Do not allow participants to attach a judgement or solution to their observation. In this work form, only what you see may be named — not yet what you think of it.",
            },
            {
                "key": "voelen",
                "label": "PHASE: FEELING",
                "name": "The Empathy Cards",
                "doel": "Aim: Broadens the emotional repertoire — which emotions might this situation evoke, even if it is not your own first reaction?",
                "cards": (
                    "Cards needed: a set of emotion cards with a single feeling each — no explanation — "
                    "such as: powerlessness, relief, shame, irritation, doubt, pride, fear, grief."
                ),
                "steps": [
                    "After hearing the case, each participant draws one emotion card blindly.",
                    "Each participant explains where or for whom in the case that specific emotion might arise — even if it is not the emotion they felt first themselves.",
                    "The group may add: for whom in the case might this feeling be even stronger?",
                    "Close with a short round: which emotion on the table had you not named yourself, but do you now recognise?",
                ],
                "tip": "Practical tip — Explicitly allow 'unpleasant' or socially inappropriate-sounding emotions (relief that it happened to someone else, for example). The aim is honest recognition, not a socially desirable feeling.",
            },
            {
                "key": "wegen",
                "label": "PHASE: WEIGHING",
                "name": "The Scale on the Floor",
                "doel": "Aim: Makes weighing arguments physical and visible, instead of remaining abstract in someone's head.",
                "cards": (
                    "Cards needed: separate argument cards (pre-printed or filled in on the spot) with short arguments for and against a possible choice in the case."
                ),
                "steps": [
                    "Mark an imaginary line on the floor or a long table, from 'fully agree' at one end to 'fully disagree' at the other.",
                    "Each participant receives one or more argument cards and places them physically on the line, based on how heavily that argument weighs for them.",
                    "Once all cards are placed, participants may question each other's placement: 'Why is your card closer to the middle than mine?'",
                    "Close with the question: if you had to merge the cards into one shared judgement now, where would the whole roughly lie?",
                ],
                "tip": "Practical tip — Deliberately allow cards to lie close to the middle. A dilemma that easily resolves to one side of the line is usually not a good dilemma.",
            },
            {
                "key": "handelen",
                "label": "PHASE: ACTING",
                "name": "The Action Sentence Set",
                "doel": "Aim: Bridges the gap between knowing what would be good and actually saying or doing it concretely — often the weakest link in moral behaviour.",
                "cards": (
                    "Cards needed: cards with the start of an action sentence, for example: "
                    "'The first thing I would do now is…', 'What I would say aloud to [name] is…', "
                    "'The first phone call I would make is to…', 'If I were standing here now, I would literally say…'"
                ),
                "steps": [
                    "After discussing the dilemma, each participant draws an action-sentence card blindly.",
                    "The participant must complete the sentence aloud — concretely and specifically (a name, a moment, a literal formulation), not in generalities such as 'I would bring it up'.",
                    "The group may probe when the answer stays too vague: 'What exactly do you say, and to whom?'",
                    "Discuss finally: which action sentence was easiest to think of, and which was hardest to actually say aloud?",
                ],
                "tip": "Practical tip — Be strict on concreteness: 'I would say something about it' is not a valid answer. Keep probing until the participant speaks a sentence they could literally use tomorrow.",
            },
        ],
        "overview_title": "Overview at a glance",
        "overview_headers": ["Phase", "Work form", "Cards needed"],
        "overview_rows": [
            ("SEEING", "The Card Fan", "Signal cards with short observation prompts"),
            ("FEELING", "The Empathy Cards", "Emotion cards — one feeling per card, no explanation"),
            ("WEIGHING", "The Scale on the Floor", "Argument cards for and against a possible choice"),
            ("ACTING", "The Action Sentence Set", "Cards with the start of an action sentence"),
        ],
        "teacher_title": "Facilitator note — safe learning environment",
        "teacher_text": (
            "Do not let participants formulate a solution or judgement during Seeing and Feeling — that comes only at Weighing and Acting. "
            "Mixing the phases is the most common pitfall: a group that jumps too quickly to 'what should we do' often skips observation and feeling, "
            "and thereby misses precisely the signals that make the dilemma hardest."
        ),
        "tip_combo": "Tip: This worksheet works well combined with conversation cards from the library — use one case card as the shared case for all four work forms.",
        "sources": (
            "Sources: Rest, J.R. (1986), Moral Development: Advances in Research and Theory  ·  "
            "Kidder, R.M. (2005), How Good People Make Tough Choices  ·  "
            "Tjin A Djie & Zwaan (2010), Managing diversity in the workplace"
        ),
        "footer": "Fontys University of Applied Sciences — Research Group Ethical Practice · Community of Moral Craftsmanship",
        "how_label": "How it works",
    },
}


def styled_run(paragraph, text, *, bold=False, italic=False, size=11, color=C_BODY):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def add_header_table(doc, lang_data):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    cells = table.rows[0].cells
    styled_run(cells[0].paragraphs[0], "Fontys Lectoraat Ethisch Werken", size=9, color=C_MUTED)
    mid = cells[1].paragraphs[0]
    for i, line in enumerate(lang_data["header_mid"].split("\n")):
        if i:
            mid.add_run("\n")
        styled_run(mid, line, bold=(i == 0), size=10 if i else 11, color=C_TITLE)
    styled_run(cells[2].paragraphs[0], lang_data["header_right"], size=9, color=C_MUTED)
    doc.add_paragraph()


def add_meta_table(doc, lang):
    labels = (
        ["Naam", "Datum", "Opleiding", "Klas / Groep"]
        if lang == "nl"
        else ["Name", "Date", "Programme", "Class / Group"]
    )
    table = doc.add_table(rows=1, cols=4)
    for cell, label in zip(table.rows[0].cells, labels):
        styled_run(cell.paragraphs[0], label, bold=True, size=10, color=C_TITLE)
    doc.add_paragraph()


def add_phase_section(doc, phase, how_label):
    color = PHASE_COLORS[phase["key"]]
    p = doc.add_paragraph()
    styled_run(p, phase["label"], bold=True, size=12, color=color)
    p = doc.add_paragraph()
    styled_run(p, phase["name"], bold=True, size=14, color=C_TITLE)
    doc.add_paragraph()
    styled_run(doc.add_paragraph(), phase["doel"], size=11, color=C_BODY)
    doc.add_paragraph()
    styled_run(doc.add_paragraph(), phase["cards"], size=11, color=C_BODY)
    doc.add_paragraph()
    styled_run(doc.add_paragraph(), how_label, bold=True, size=11, color=C_TITLE)
    for i, step in enumerate(phase["steps"], 1):
        styled_run(doc.add_paragraph(), f"{i}. {step}", size=11, color=C_BODY)
    doc.add_paragraph()
    tip = doc.add_paragraph()
    styled_run(tip, phase["tip"], italic=True, size=10, color=color)
    doc.add_paragraph()


def build_doc(lang):
    data = CONTENT[lang]
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_header_table(doc, data)
    add_meta_table(doc, lang)

    p = doc.add_paragraph()
    styled_run(p, data["title"], bold=True, size=22, color=C_TITLE)
    styled_run(doc.add_paragraph(), data["subtitle"], italic=True, size=11, color=C_THEMES)
    doc.add_paragraph()

    styled_run(doc.add_paragraph(), data["intro_title"], bold=True, size=13, color=C_TITLE)
    styled_run(doc.add_paragraph(), data["intro_p1"], size=11, color=C_BODY)
    styled_run(doc.add_paragraph(), data["intro_p2"], size=11, color=C_BODY)
    doc.add_paragraph()

    for phase in data["phases"]:
        add_phase_section(doc, phase, data["how_label"])

    styled_run(doc.add_paragraph(), data["overview_title"], bold=True, size=13, color=C_TITLE)
    table = doc.add_table(rows=1 + len(data["overview_rows"]), cols=3)
    table.style = "Table Grid"
    for j, header in enumerate(data["overview_headers"]):
        styled_run(table.rows[0].cells[j].paragraphs[0], header, bold=True, size=10, color=C_TITLE)
    for ri, row in enumerate(data["overview_rows"], 1):
        for ci, cell_text in enumerate(row):
            styled_run(table.rows[ri].cells[ci].paragraphs[0], cell_text, size=10, color=C_BODY)
    doc.add_paragraph()

    styled_run(doc.add_paragraph(), data["teacher_title"], bold=True, size=12, color=C_TITLE)
    styled_run(doc.add_paragraph(), data["teacher_text"], size=11, color=C_BODY)
    doc.add_paragraph()

    combo = doc.add_paragraph()
    styled_run(combo, data["tip_combo"], italic=True, size=10, color=RGBColor(0x53, 0x4A, 0xB7))
    doc.add_paragraph()

    src = doc.add_paragraph()
    styled_run(src, data["sources"], size=9, color=C_MUTED)
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(foot, data["footer"], italic=True, size=8, color=C_MUTED)

    out = OUT_DIR / data["filename"]
    doc.save(out)
    return out


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for lang in ("nl", "en"):
        path = build_doc(lang)
        print(path.relative_to(ROOT))


if __name__ == "__main__":
    main()

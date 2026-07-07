"""Genereer losse .docx per onderwijs-kaart (NL + EN) — v3 layout, A5, categoriekleur."""
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CARDS_PATH = ROOT / "src/data/gesprekskaarten/cards.json"
OUT_DIR = ROOT / "public/downloads/gesprekskaarten/onderwijs"
SET_ID = "onderwijs-v3"

COMPLEXITY_LABELS = {
    "nl": {
        "micro": "★☆☆ Micro — persoonlijke keuze",
        "meso": "★★☆ Meso — organisatorisch dilemma",
        "macro": "★★★ Macro — systemisch dilemma",
    },
    "en": {
        "micro": "★☆☆ Micro — direct, personal choice",
        "meso": "★★☆ Meso — organisational dilemma",
        "macro": "★★★ Macro — systemic dilemma",
    },
}

CATEGORY_LABELS = {
    "nl": {"onderwijs": "Onderwijs"},
    "en": {"onderwijs": "Education"},
}

CATEGORY_COLORS = {
    "onderwijs": RGBColor(0x1A, 0x90, 0x80),
}

COMPLEXITY_COLOR = RGBColor(0x99, 0x35, 0x56)
ATTRIBUTION_COLOR = RGBColor(0x5F, 0x5E, 0x5A)

ATTRIBUTION = {
    "nl": 'Complexiteitsmodel: Kim Meijer, "Impossible and Inevitable" (Tilburg University)',
    "en": 'Complexity model: Kim Meijer, "Impossible and Inevitable" (Tilburg University)',
}


def configure_a5(doc):
    section = doc.sections[0]
    section.page_height = Mm(210)
    section.page_width = Mm(148)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)


def write_card_doc(card, lang):
    content = card[lang]
    doc = Document()
    configure_a5(doc)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    cat = CATEGORY_LABELS[lang].get(card["categorie"], card["categorie"])
    cat_color = CATEGORY_COLORS.get(card["categorie"], ATTRIBUTION_COLOR)
    complexity = COMPLEXITY_LABELS[lang][card["complexiteit"]]

    p = doc.add_paragraph(cat.upper())
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = cat_color

    p2 = doc.add_paragraph(complexity)
    p2.runs[0].bold = True
    p2.runs[0].font.color.rgb = COMPLEXITY_COLOR

    title = doc.add_paragraph(content["titel"])
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x27, 0x44)

    story = doc.add_paragraph(content["verhaal"])
    story.paragraph_format.space_after = Pt(12)
    for run in story.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    q = doc.add_paragraph()
    q.add_run(content.get("vraag1", "")).bold = True
    q2 = doc.add_paragraph()
    q2.add_run(content.get("vraag2", "")).bold = True

    attr = doc.add_paragraph(ATTRIBUTION[lang])
    attr.runs[0].italic = True
    attr.runs[0].font.size = Pt(8)
    attr.runs[0].font.color.rgb = ATTRIBUTION_COLOR

    suffix = "NL" if lang == "nl" else "EN"
    out = OUT_DIR / f"{card['id']}_{suffix}.docx"
    doc.save(out)
    return out


def main():
    cards = json.loads(CARDS_PATH.read_text(encoding="utf-8"))
    subset = [c for c in cards if c.get("set") == SET_ID]
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    for card in subset:
        for lang in ("nl", "en"):
            path = write_card_doc(card, lang)
            print(path.relative_to(ROOT))

    print(f"Klaar: {len(subset) * 2} bestanden in {OUT_DIR.relative_to(ROOT)}")


if __name__ == "__main__":
    main()

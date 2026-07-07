"""Genereer losse .docx per burgerschapskaart (NL + EN)."""
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
CARDS_PATH = ROOT / "src/data/gesprekskaarten/cards.json"
OUT_DIR = ROOT / "public/downloads/gesprekskaarten/burgerschap"

COMPLEXITY_LABELS = {
    "nl": {
        "micro": "★☆☆ Micro — persoonlijke keuze",
        "meso": "★★☆ Meso — organisatorisch dilemma",
        "macro": "★★★ Macro — systemisch dilemma",
    },
    "en": {
        "micro": "★☆☆ Micro — direct, personal choice",
        "meso": "★★☆ Meso — organisational dilemma",
        "macro": "★★★ Macro — systemic dilemma",
    },
}

CATEGORY_LABELS = {
    "nl": {
        "werk": "Werk",
        "dagelijks-leven": "Dagelijks leven",
        "social-media": "Social media",
        "studentenleven": "Studentenleven",
        "diversiteit-inclusie": "Diversiteit & inclusie",
        "duurzaamheid": "Duurzaamheid",
        "onderwijs": "Onderwijs",
    },
    "en": {
        "werk": "Work",
        "dagelijks-leven": "Daily life",
        "social-media": "Social media",
        "studentenleven": "Student life",
        "diversiteit-inclusie": "Diversity & inclusion",
        "duurzaamheid": "Sustainability",
        "onderwijs": "Education",
    },
}

ATTRIBUTION = {
    "nl": 'Complexiteitsmodel: Kim Meijer, "Impossible and Inevitable" (Tilburg University)',
    "en": 'Complexity model: Kim Meijer, "Impossible and Inevitable" (Tilburg University)',
}


def write_card_doc(card, lang):
    content = card[lang]
    doc = Document()
    cat = CATEGORY_LABELS[lang].get(card["categorie"], card["categorie"])
    complexity = COMPLEXITY_LABELS[lang][card["complexiteit"]]

    p = doc.add_paragraph(cat.upper())
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(0x5F, 0x5E, 0x5A)

    p2 = doc.add_paragraph(complexity)
    p2.runs[0].bold = True
    p2.runs[0].font.color.rgb = RGBColor(0x99, 0x35, 0x56)

    title = doc.add_paragraph(content["titel"])
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    story = doc.add_paragraph(content["verhaal"])
    story.paragraph_format.space_after = Pt(12)

    q = doc.add_paragraph()
    q.add_run(content.get("vraag1", "")).bold = True
    q2 = doc.add_paragraph()
    q2.add_run(content.get("vraag2", "")).bold = True

    attr = doc.add_paragraph(ATTRIBUTION[lang])
    attr.runs[0].italic = True
    attr.runs[0].font.size = Pt(8)
    attr.runs[0].font.color.rgb = RGBColor(0x5F, 0x5E, 0x5A)

    suffix = "NL" if lang == "nl" else "EN"
    out = OUT_DIR / f"{card['id']}_{suffix}.docx"
    doc.save(out)
    return out


def main():
    cards = json.loads(CARDS_PATH.read_text(encoding="utf-8"))
    burgerschap = [c for c in cards if c.get("set") == "burgerschap"]
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    for card in burgerschap:
        for lang in ("nl", "en"):
            path = write_card_doc(card, lang)
            print(path.relative_to(ROOT))

    print(f"Klaar: {len(burgerschap) * 2} bestanden in {OUT_DIR.relative_to(ROOT)}")


if __name__ == "__main__":
    main()

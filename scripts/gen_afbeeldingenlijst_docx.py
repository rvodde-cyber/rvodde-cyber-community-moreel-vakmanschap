# -*- coding: utf-8 -*-
"""Genereer afbeeldingenlijst voor filosofie/ethiek manuscript."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = (
    r"C:\Users\876409\OneDrive - Office 365 Fontys\Documenten"
    r"\Ethiek en filpsofie lessenreeks\Afbeeldingenlijst_filosofie_en_ethiek.docx"
)

AFBEELDINGEN = [
  {
    "nr": 1,
    "titel": "Tijdlijn westerse filosofie",
    "waar": "Inleiding — na titel",
    "zoeken": "timeline western philosophy infographic; geschiedenis filosofie tijdlijn",
    "firefly": (
      "Horizontal educational timeline from ancient Greece 500 BC to present day, "
      "five labeled eras: Classical Antiquity, Middle Ages, Enlightenment, Modern Era, "
      "Digital Age, small iconic symbols per era, minimalist infographic, soft blue and "
      "sand colors, white background, no readable text"
    ),
    "tip": "Infographic werkt beter in Canva/PowerPoint als je Nederlandse labels wilt.",
  },
  {
    "nr": 2,
    "titel": "Kaart van de Griekse wereld",
    "waar": "Voor Deel 1 — Klassieke oudheid",
    "zoeken": "ancient Greece map Athens Sparta; Mediterranean map classical Greece",
    "firefly": (
      "Vintage-style educational map of ancient Greece and Mediterranean, Athens and Sparta "
      "marked, muted parchment tones, clean markers without text, textbook illustration"
    ),
    "tip": "Wikimedia heeft goede historische kaarten (publiek domein).",
  },
  {
    "nr": 3,
    "titel": "De Agora en de geboorte van filosofie",
    "waar": "Deel 1 — tijdsbeeld",
    "zoeken": "Ancient Agora Athens; Greek philosophers discussion agora",
    "firefly": (
      "Split illustration: left side ancient Greek agora with citizens in dialogue, "
      "right side abstract mythological lightning versus rational geometric symbols, "
      "educational contrast, warm Mediterranean light, painterly but clean"
    ),
    "tip": "Foto van de Agora van Athene (Unsplash/Wikimedia) + apart symboolbeeld voor mythe vs. logos.",
  },
  {
    "nr": 4,
    "titel": "Socrates en de hemlock",
    "waar": "Na paragraaf Socrates",
    "zoeken": "Socrates hemlock David painting; Socrates bust Vatican Wikimedia",
    "firefly": (
      "Classical painting style scene inspired by ancient Greek philosophy, elderly philosopher "
      "calmly accepting a cup in a dignified setting, restrained dramatic lighting, historical "
      "fine art style, respectful tone, no gore"
    ),
    "tip": "Liever echte buste of schilderij David via Wikimedia dan AI-portret.",
  },
  {
    "nr": 5,
    "titel": "Plato's allegorie van de grot",
    "waar": "Na paragraaf Plato",
    "zoeken": "Plato cave allegory diagram; allegory of the cave illustration",
    "firefly": (
      "Cross-section diagram of Plato's cave allegory: prisoners watching shadows on wall, "
      "fire behind them, one figure climbing toward bright opening, clear educational diagram "
      "style, dark cave versus golden light outside, no text labels"
    ),
    "tip": "Maak ook een moderne variant: schermen als schaduwen op de muur.",
  },
  {
    "nr": "5b",
    "titel": "Plato's grot — moderne variant (optioneel)",
    "waar": "Na paragraaf Plato — als tweede beeld",
    "zoeken": "social media cave allegory modern; filter bubble illustration",
    "firefly": (
      "Modern allegory of the cave: people staring at smartphone screens showing shadows, "
      "one person walking toward real daylight outside, subtle critique of social media, "
      "educational illustration"
    ),
    "tip": "Sluit aan bij de moderne vertaalslag in het manuscript.",
  },
  {
    "nr": 6,
    "titel": "Aristoteles — deugd als middenweg",
    "waar": "Na paragraaf Aristoteles",
    "zoeken": "golden mean Aristotle virtue ethics diagram; Aristotle bust",
    "firefly": (
      "Educational spectrum diagram showing virtue as balance between two extremes, example "
      "courage between cowardice and recklessness, balanced scale in center, minimalist "
      "infographic, sage green and cream colors, no text"
    ),
    "tip": "Diagram met Nederlandse termen het beste in Canva.",
  },
  {
    "nr": 7,
    "titel": "Middeleeuwse scholastiek",
    "waar": "Deel 2 — tijdsbeeld",
    "zoeken": "medieval scriptorium monks; medieval university Paris lecture",
    "firefly": (
      "Medieval European scriptorium, scholar copying manuscripts by candlelight, stained "
      "glass light, calm scholarly atmosphere, historical educational illustration, warm amber tones"
    ),
    "tip": "Alternatief: schema 'Geloof + Rede → synthese' in Canva.",
  },
  {
    "nr": 8,
    "titel": "Augustinus",
    "waar": "Na paragraaf Augustinus",
    "zoeken": "Saint Augustine portrait; Augustine of Hippo mosaic",
    "firefly": (
      "Byzantine-style portrait of early Christian bishop philosopher in North Africa, "
      "contemplative expression, gold and deep blue tones, historical religious art style, dignified"
    ),
    "tip": "Wikimedia: mozaïek of schilderij Augustinus.",
  },
  {
    "nr": 9,
    "titel": "Thomas van Aquino",
    "waar": "Na paragraaf Aquino",
    "zoeken": "Thomas Aquinas painting sun on chest; natural law ethics diagram",
    "firefly": (
      "Medieval scholastic philosopher in Dominican robes with subtle radiant symbol on chest, "
      "surrounded by open books and geometric order, synthesis of faith and reason theme, "
      "Renaissance painting style"
    ),
    "tip": "Diagram natuurlijke + theologische deugden als Venn-diagram in Canva.",
  },
  {
    "nr": 10,
    "titel": "De Verlichting",
    "waar": "Deel 3 — tijdsbeeld",
    "zoeken": "Enlightenment salon 18th century; Age of Enlightenment illustration",
    "firefly": (
      "18th century Enlightenment salon, diverse thinkers in discussion, globe and books, "
      "soft daylight through tall windows, elegant educational historical scene, optimism and reason theme"
    ),
    "tip": "Combineer met portretten Descartes, Kant, Mill (Wikimedia).",
  },
  {
    "nr": 11,
    "titel": "Descartes (aanbevolen)",
    "waar": "Bij paragraaf Descartes",
    "zoeken": "René Descartes portrait; cogito ergo sum illustration",
    "firefly": (
      "Portrait of 17th century French philosopher at writing desk with geometric diagrams "
      "and candle, atmosphere of doubt and discovery, Dutch Golden Age painting style"
    ),
    "tip": "Echt portret via Wikimedia/Rijksmuseum heeft voorkeur.",
  },
  {
    "nr": 12,
    "titel": "Kant en de categorische imperatief",
    "waar": "Na paragraaf Kant",
    "zoeken": "Immanuel Kant portrait; categorical imperative ethics flowchart",
    "firefly": (
      "Educational flowchart concept: moral decision path from personal rule to universal test, "
      "fork in road yes/no, human dignity icon, minimalist blue-gray infographic, no readable text"
    ),
    "tip": "Flowchart: Maxime → universaliseerbaar? → Ja/Nee in Canva met NL-tekst.",
  },
  {
    "nr": 13,
    "titel": "Mill en het harmprincipe",
    "waar": "Na paragraaf Mill",
    "zoeken": "harm principle liberty illustration; John Stuart Mill portrait",
    "firefly": (
      "Educational diagram: large circle as personal freedom zone, outer boundary where harm "
      "to others begins, simple clean infographic, subtle industrial era background, liberal "
      "ethics theme, no text"
    ),
    "tip": "Portret Mill via Wikimedia; diagram in Canva.",
  },
  {
    "nr": 14,
    "titel": "Moderne tijd en existentiële crisis",
    "waar": "Deel 4 — tijdsbeeld",
    "zoeken": "industrial revolution factory smoke; WWI battlefield horizon; Holocaust memorial abstract",
    "firefly": (
      "Respectful educational collage: industrial factory smoke, empty battlefield horizon, "
      "abstract memorial geometry, muted gray and sepia palette, somber but not graphic, "
      "historical reflection theme"
    ),
    "tip": "Holocaust/WOII: gebruik echte monumentfoto's met bronvermelding, geen AI.",
  },
  {
    "nr": 15,
    "titel": "Nietzsche (aanbevolen)",
    "waar": "Bij paragraaf Nietzsche",
    "zoeken": "Friedrich Nietzsche portrait; philosophy values crisis mountain",
    "firefly": (
      "19th century German philosopher portrait with intense gaze, stormy mountain landscape "
      "behind symbolizing crisis of values, dramatic but thoughtful, not propaganda style"
    ),
    "tip": "Vermijd AI die naar nazi-propaganda neigt; gebruik historisch portret.",
  },
  {
    "nr": 16,
    "titel": "Hannah Arendt en de publieke ruimte",
    "waar": "Na paragraaf Arendt",
    "zoeken": "Hannah Arendt portrait; public space agora democracy",
    "firefly": (
      "Split scene: vibrant public town square with people debating versus isolated individuals "
      "in separate glowing bubbles, political philosophy theme, clean editorial illustration"
    ),
    "tip": "Portret Arendt via Wikimedia; geen sensatiebeelden Eichmann-proces.",
  },
  {
    "nr": 17,
    "titel": "Sartre en Beauvoir (aanbevolen)",
    "waar": "Bij paragraaf Sartre & Beauvoir",
    "zoeken": "Sartre Beauvoir Paris cafe; existentialism freedom illustration",
    "firefly": (
      "1940s Paris cafe scene, two intellectuals in conversation, thoughtful atmosphere, "
      "black and white photographic style, authentic and dignified"
    ),
    "tip": "Historische foto's vaak beschikbaar via Wikimedia.",
  },
  {
    "nr": 18,
    "titel": "Rawls en het sluier van onwetendheid",
    "waar": "Na paragraaf Rawls",
    "zoeken": "veil of ignorance illustration; John Rawls justice",
    "firefly": (
      "Symbolic illustration: diverse silhouettes standing behind a translucent veil curtain, "
      "equal footing, designing society together, soft justice theme colors, educational, no text"
    ),
    "tip": "Sluier-metafoor werkt goed als eenvoudige symbolische illustratie.",
  },
  {
    "nr": 19,
    "titel": "Camus (aanbevolen)",
    "waar": "Bij paragraaf Camus",
    "zoeken": "Sisyphus myth illustration; Albert Camus portrait",
    "firefly": (
      "Myth of Sisyphus: figure pushing boulder up hill at dawn, absurd yet dignified, "
      "Mediterranean landscape, existential philosophy theme, painterly minimalist style"
    ),
    "tip": "Sisyphus is een sterk symbolisch beeld; geen grappige memes.",
  },
  {
    "nr": 20,
    "titel": "De digitale agora",
    "waar": "Deel 5 — begin",
    "zoeken": "digital public sphere illustration; social media agora metaphor",
    "firefly": (
      "Modern town square formed by floating smartphones and laptops, people connecting through "
      "light streams, subtle Plato cave shadows on screens in background, hopeful tech-society "
      "illustration, clean and contemporary"
    ),
    "tip": "Niet te technofobisch; educatief en uitnodigend.",
  },
  {
    "nr": 21,
    "titel": "Casus AI in de zorg",
    "waar": "Deel 5 — groepswerk-casus",
    "zoeken": "AI healthcare ethics stakeholders diagram; healthcare AI infographic",
    "firefly": (
      "Circular stakeholder diagram around AI brain icon in hospital setting: patient, nurse, "
      "doctor, IT vendor, government regulator, connected by lines, neutral healthcare colors, "
      "educational infographic, no brand logos, no text"
    ),
    "tip": "Stakeholderkaart het beste in Canva met NL-labels.",
  },
  {
    "nr": 22,
    "titel": "Vergelijkende matrix",
    "waar": "Slot — overzicht 'het goede'",
    "zoeken": "ethics comparison chart virtue deontology utilitarianism",
    "firefly": (
      "Clean comparison matrix grid with color-coded rows for historical eras, columns for ethics "
      "themes, poster-style educational design, pastel columns, no readable text"
    ),
    "tip": "Maak in Excel/Canva met denkernamen en thema's ingevuld.",
  },
  {
    "nr": 23,
    "titel": "Kaart aanvullende stromingen",
    "waar": "Aanvullend hoofdstuk stromingen",
    "zoeken": "philosophy mind map contemporary ethics",
    "firefly": (
      "Mind map cluster diagram with center node present day, branches to pragmatism phenomenology "
      "critical theory care ethics communitarianism decolonial thought technology ethics climate "
      "ethics, modern infographic, soft colors, no text"
    ),
    "tip": "Labels in Canva op Nederlandse termen zetten.",
  },
]

STROMINGEN_EXTRA = [
  ("Pragmatisme", "John Dewey education experiential learning",
   "Teacher and student learning by doing in workshop, pragmatist philosophy vibe, warm natural light"),
  ("Fenomenologie", "lived experience healthcare patient perspective",
   "First-person view of caring hands meeting patient, embodied experience theme, soft focus, empathetic"),
  ("Kritische theorie", "culture industry mass media critique",
   "Crowd absorbed by identical glowing screens, subtle critique of mass media, editorial illustration"),
  ("Foucault / poststructuralisme", "panopticon surveillance diagram",
   "Abstract panopticon tower with observing eye, power and knowledge theme, minimalist architectural diagram"),
  ("Zorgethiek", "ethics of care relationship",
   "Two people in supportive conversation, care and trust, gentle illustration for healthcare ethics"),
  ("Communitarisme", "professional community craftsmanship",
   "Circle of craftspeople sharing tools and knowledge, community of practice theme, warm collaborative scene"),
  ("Capability approach", "human development opportunities infographic",
   "Diverse people accessing education health mobility as opportunities, development ethics infographic"),
  ("Decoloniaal denken", "colonial legacy knowledge power",
   "World map with multiple knowledge traditions as equal streams merging, decolonial education theme, respectful"),
  ("Techniekfilosofie", "AI ethics human machine relationship",
   "Human and transparent AI interface handshake, responsibility and technology theme, calm futuristic"),
  ("Klimaatethiek", "intergenerational justice climate",
   "Older and younger person looking at changing landscape together, intergenerational responsibility, hopeful tone"),
]


def add_label_value(doc, label, value, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    run = p.add_run(value)
    if italic:
        run.italic = True


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Afbeeldingenlijst", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Bijlage bij: Een geschiedenis van de filosofie en ethiek"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Zoektermen voor beeldbanken (Wikimedia, Unsplash) en prompts voor Adobe Firefly. "
        "Doelgroep: HBO — stijl: strak, educatief, veel witruimte."
    )

    doc.add_heading("Algemene Firefly-suffix", 1)
    doc.add_paragraph(
        "Voeg aan elke prompt toe voor consistente stijl:"
    )
    p = doc.add_paragraph()
    p.add_run(
        ", educational illustration, clean layout, soft muted colors, white background, "
        "professional textbook style, no text, no watermark"
    ).italic = True

    doc.add_heading("Portretten: zoeken vs. genereren", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Denker"
    hdr[1].text = "Aanbevolen bron"
    hdr[2].text = "Firefly?"
    rows = [
      ("Socrates, Plato, Aristoteles", "Wikimedia bustes", "Liever echt beeld"),
      ("Kant, Mill, Descartes", "Wikimedia / Rijksmuseum", "Liever echt beeld"),
      ("Arendt, Beauvoir, Rawls", "Wikimedia", "Liever echt beeld"),
      ("Nietzsche, Camus", "Wikimedia", "Liever echt beeld"),
      ("Augustinus, Aquino", "Religieuze kunst Wikimedia", "Stijl OK in Firefly"),
      ("Diagrammen / infographics", "Canva / PowerPoint", "Firefly beperkt voor tekst"),
    ]
    for denker, bron, firefly in rows:
        row = table.add_row().cells
        row[0].text = denker
        row[1].text = bron
        row[2].text = firefly

    doc.add_heading("Afbeeldingen per hoofdstuk", 1)

    for item in AFBEELDINGEN:
        nr = item["nr"]
        doc.add_heading(f"#{nr} — {item['titel']}", 2)
        add_label_value(doc, "Plaats in manuscript", item["waar"])
        add_label_value(doc, "Zoektermen", item["zoeken"])
        add_label_value(doc, "Adobe Firefly-prompt", item["firefly"], italic=True)
        add_label_value(doc, "Tip", item["tip"])

    doc.add_heading("Extra — aanvullende stromingen (optioneel)", 1)
    doc.add_paragraph(
        "Eén beeld per stroming uit het aanvullende hoofdstuk, indien je dat visueel wilt maken."
    )
    for naam, zoeken, firefly in STROMINGEN_EXTRA:
        doc.add_heading(naam, 3)
        add_label_value(doc, "Zoektermen", zoeken)
        add_label_value(doc, "Firefly-prompt", firefly, italic=True)

    doc.add_heading("Workflow", 1)
    for stap in [
        "Historische portretten ophalen via Wikimedia Commons (controleer licentie).",
        "Symbolische illustraties genereren in Adobe Firefly met prompts uit deze lijst.",
        "Infographics met Nederlandse tekst maken in Canva of PowerPoint.",
        "In Word: [AFBEELDING]-placeholder vervangen → Invoegen → Afbeelding.",
        "Bronvermelding onder elke afbeelding in het manuscript zetten.",
    ]:
        doc.add_paragraph(stap, style="List Number")

    doc.save(OUT)
    print(f"Opgeslagen: {OUT}")


if __name__ == "__main__":
    main()

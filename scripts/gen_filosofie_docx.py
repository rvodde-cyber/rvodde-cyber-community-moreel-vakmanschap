# -*- coding: utf-8 -*-
"""Genereer uitgebreid filosofie/ethiek manuscript voor HBO."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = (
    r"C:\Users\876409\OneDrive - Office 365 Fontys\Documenten"
    r"\Ethiek en filpsofie lessenreeks\Een geschiedenis van de filosofie en ethiek - redactie.docx"
)


def add_image_placeholder(doc, title, description):
    p = doc.add_paragraph()
    run = p.add_run(f"[AFBEELDING: {title}]")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x7A)
    doc.add_paragraph(description).italic = True


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    doc.add_paragraph(text)


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def add_verwerkingsvraag(doc, denker, casus, vraag):
    add_subheading(doc, f"Verwerkingsvraag — vanuit {denker}")
    p = doc.add_paragraph()
    r1 = p.add_run("Casus: ")
    r1.bold = True
    p.add_run(casus)
    p = doc.add_paragraph()
    r2 = p.add_run("Wat zou je vanuit het denken van ")
    r2.bold = True
    p.add_run(f"{denker} nu doen? {vraag}")


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Titel
    title = doc.add_heading("Het Boek van de Westerse Geest: Van Verwondering tot Rebellie", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(
        doc,
        "Een overzicht van de westerse filosofie en ethiek voor HBO-studenten. "
        "Dit document helpt je de grote denkers en hun ideeën te begrijpen — "
        "niet als abstracte theorie, maar als antwoorden op vragen die ook vandaag "
        "in je beroepspraktijk spelen: Wat is eerlijk? Wat is mijn verantwoordelijkheid? "
        "Hoe ga ik om met macht, technologie en de ander?",
    )

    add_image_placeholder(
        doc,
        "Tijdlijn westerse filosofie",
        "Horizontale tijdlijn van ca. 500 v.Chr. tot heden met de vijf hoofdperiodes "
        "(Oudheid, Middeleeuwen, Verlichting, Moderne tijd, Digitale tijd). "
        "Per periode 2–3 portretten of iconen. Gebruik een rustige, leesbare infographic.",
    )

    # INLEIDING
    add_heading(doc, "Inleiding", 1)
    add_body(
        doc,
        "Filosofie betekent letterlijk 'liefde voor wijsheid'. Het is geen verzameling "
        "feiten, maar een manier van nadenken over de fundamenten onder ons bestaan: "
        "kennis, moraal, vrijheid, rechtvaardigheid en zin. Ethiek is het deelgebied "
        "dat zich richt op de vraag: wat moet ik doen — en waarom?",
    )
    add_body(
        doc,
        "Dit manuscript volgt de rode draad van het westerse denken. We kijken telkens "
        "naar drie dingen: (1) de historische context — welke wereld vormde de denker? "
        "(2) de kernideeën — wat zei hij of zij over 'het goede'? en (3) de moderne "
        "vertaalslag — wat betekent dit voor jouw werk en samenleving?",
    )
    add_subheading(doc, "Hoe gebruik je dit document?")
    add_bullet(doc, "Lees per hoofdstuk eerst het tijdsbeeld: context maakt ideeën begrijpelijk.")
    add_bullet(doc, "Let op de kaders 'Het Goede' en 'Ethische gedachte': dat is de kern per denker.")
    add_bullet(doc, "Gebruik de verwerkingsvragen per denker: casus + 'Wat zou je vanuit het denken van… nu doen?'")
    add_bullet(doc, "Gebruik de reflectievragen aan het eind van elk deel voor groepswerk of portfolio.")
    add_bullet(doc, "Let op [CHECK]-markeringen: plekken waar de tekst bewust is versimpeld voor HBO.")

    add_image_placeholder(
        doc,
        "Kaart van de Griekse wereld",
        "Kaart van het Middellandse Zeegebied met Athene, Sparta en belangrijke filosofenscholen. "
        "Optioneel: foto van de Agora van Athene als sfeerbeeld.",
    )

    # DEEL 1
    add_heading(doc, "Deel 1: De Klassieke Oudheid – De geboorte van de rede", 1)

    add_subheading(doc, "Tijdsbeeld")
    add_body(
        doc,
        "Tussen circa 600 en 300 v.Chr. bloeien de Griekse stadstaten (polis: zelfstandige "
        "stadsstaat), met Athene als belangrijk centrum. Op de agora — de markt en "
        "ontmoetingsplek — bespreken burgers politiek, handel en recht.",
    )
    add_body(
        doc,
        "Tot dan toe werd de wereld vooral verklaard vanuit mythologie: goden als oorzaak "
        "van donder, ziekte en oorlog. In deze periode verschuift het denken. "
        "Natuurverschijnselen worden steeds vaker verklaard met behulp van rede "
        "(logos: rationeel denken en argumentatie). Filosofie ontstaat als kritisch en "
        "systematisch nadenken over de werkelijkheid en het goede leven.",
    )
    add_body(
        doc,
        "Voor HBO-studenten markeert dit het begin van 'professioneel redeneren': niet "
        "zomaar aannemen wat waar is, maar overtuigingen onderzoeken en onderbouwen. "
        "Dat zie je terug in de zorg (evidence-based werken), in IT (logische systemen) "
        "en in het onderwijs (kritisch denken).",
    )

    add_image_placeholder(
        doc,
        "De Agora en de geboorte van filosofie",
        "Illustratie of foto: Griekse agora met burgers in gesprek; naastbeeld: contrast "
        "tussen mythologische voorstelling (bijv. Zeus) en rationele verklaring (bijv. "
        "Pythagoras of een vroeg natuurfilosoof).",
    )

    # Socrates
    add_heading(doc, "Socrates (ca. 470–399 v.Chr.)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Socrates schreef zelf niets op. We kennen hem via zijn leerlingen, vooral Plato. "
        "Hij was de zoon van een steenhouwer en diende als soldaat in de Peloponnesische "
        "Oorlog tussen Athene en Sparta.",
    )
    add_body(
        doc,
        "In een tijd van politieke spanning daagde hij Atheners uit hun overtuigingen te "
        "onderzoeken. Zijn uitspraak 'Ik weet dat ik niets weet' is geen vorm van "
        "bescheidenheid, maar een uitgangspunt: wie de grenzen van zijn kennis erkent, "
        "kan blijven leren.",
    )
    add_body(
        doc,
        "Athene veroordeelde hem uiteindelijk tot de dood door het drinken van hemlock "
        "(gif). Hij werd beschuldigd van het misleiden van de jeugd en goddeloosheid. "
        "Zijn leven laat zien dat kritisch denken spanning kan veroorzaken met machthebbers.",
    )
    add_subheading(doc, "Kernbegrip: de Socratische methode")
    add_body(
        doc,
        "Socrates gaf geen lezingen, maar stelde vragen. Door steeds door te vragen "
        "('Wat bedoel je precies?') bracht hij gesprekspartners tot tegenstrijdigheden "
        "in hun eigen denken.",
    )
    add_body(
        doc,
        "Deze methode wordt de elenchus genoemd (een vorm van kritisch onderzoek via "
        "dialoog). Het doel is niet winnen, maar inzicht.",
    )
    add_subheading(doc, "Het Goede")
    add_body(
        doc,
        "Voor Socrates is het goede verbonden met kennis (intellectualisme: het idee dat "
        "juist inzicht leidt tot juist handelen). Wie het goede werkelijk begrijpt, "
        "zal daarnaar handelen.",
    )
    add_body(
        doc,
        "Kwaad handelen komt volgens hem voort uit onwetendheid — niet uit een bewuste "
        "keuze voor het verkeerde.",
    )
    add_body(doc, "[CHECK: deze opvatting wordt vaak gezien als te simplistisch]")
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "De Socratische methode kan helpen om polarisatie te verminderen, bijvoorbeeld op "
        "sociale media. In plaats van snel te reageren, stel je vragen en onderzoek je "
        "aannames. Ook in teams en klantgesprekken is dit waardevol: eerst begrijpen, "
        "dan pas oplossingen formuleren.",
    )
    add_verwerkingsvraag(
        doc,
        "Socrates",
        "Op social media gaat een bericht viral waarin een medestudent wordt beschuldigd van "
        "plagiaat. Binnen een uur heeft iedereen een mening en de docent wordt onder druk "
        "gezet om direct op te treden.",
        "Welke vragen stel je voordat je een oordeel vormt? Aan wie stel je deze vragen, "
        "en wat hoop je daarmee te verduidelijken?",
    )

    add_image_placeholder(
        doc,
        "Socrates en de hemlock",
        "Klassiek schilderij van Socrates die de hemlockbeker aanvaardt (bijv. David), "
        "of een portretbuste. Bijschrift: 'De filosoof die liever stierf dan stopte met vragen stellen.'",
    )

    # Plato
    add_heading(doc, "Plato (ca. 427–347 v.Chr.)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Plato was aristocraat en leerling van Socrates. De dood van zijn leermeester "
        "schokte hem diep. Hij vroeg zich af hoe een democratie zo'n beslissing kon nemen "
        "en ging op zoek naar zekere kennis.",
    )
    add_body(
        doc,
        "Zijn antwoord lag in de theorie van de Ideeën (Formen: onveranderlijke en perfecte "
        "modellen van werkelijkheid). Volgens Plato is de wereld die we waarnemen slechts "
        "een afgeleide van deze diepere werkelijkheid.",
    )
    add_subheading(doc, "Het Goede")
    add_body(
        doc,
        "Het Goede is voor Plato de hoogste Idee. In zijn filosofie vervult het een "
        "vergelijkbare rol als de zon: het maakt inzicht mogelijk en geeft richting aan "
        "andere waarden, zoals rechtvaardigheid en wijsheid.",
    )
    add_subheading(doc, "De grot — uitgelegd")
    add_body(
        doc,
        "In de allegorie van de grot beschrijft Plato mensen die vastzitten en alleen "
        "schaduwen zien op een muur. Ze denken dat dit de werkelijkheid is.",
    )
    add_body(
        doc,
        "Wanneer iemand vrijkomt en naar buiten gaat, ontdekt hij de echte wereld. Keert "
        "hij terug, dan wordt hij niet geloofd. Plato laat hiermee zien dat mensen vaak "
        "vasthouden aan beperkte kennis en moeite hebben met nieuwe inzichten.",
    )
    add_body(doc, "[CHECK: interpretatie 'de meeste mensen leven in schijn' is didactisch versimpeld]")
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "In een digitale wereld lijkt Plato verrassend actueel. Algoritmes bepalen welke "
        "informatie je ziet, waardoor je een gefilterd wereldbeeld krijgt. Professionals "
        "moeten zich bewust zijn van deze 'schaduwen' en actief andere perspectieven zoeken.",
    )
    add_verwerkingsvraag(
        doc,
        "Plato",
        "Je merkt dat je social media vooral berichten toont die jouw mening bevestigen. "
        "Anderen zien compleet andere 'feiten'.",
        "Wat is in deze situatie schijn en wat is werkelijkheid? Welke stap zou de "
        "'bevrijde gevangene' nemen?",
    )

    add_image_placeholder(
        doc,
        "Plato's allegorie van de grot",
        "Diagram of illustratie van de grot: gevangenen, vuur, schaduwen op de muur, "
        "en de bevrijde figuur die naar buiten klimt. Eventueel moderne variant: "
        "mensen die naar smartphones kijken terwijl 'echte' werkelijkheid buiten beeld valt.",
    )

    # Aristoteles
    add_heading(doc, "Aristoteles (384–322 v.Chr.)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Aristoteles was leerling van Plato, maar legde meer nadruk op de waarneembare "
        "werkelijkheid. Hij bestudeerde systematisch onderwerpen als biologie, logica en "
        "politiek, en richtte het Lyceum op — een centrum voor onderzoek en onderwijs.",
    )
    add_subheading(doc, "Het Goede: Eudaimonia")
    add_body(
        doc,
        "Eudaimonia wordt vaak vertaald als 'geluk', maar betekent beter: menselijke bloei "
        "— een geslaagd leven waarin je je vermogens ontwikkelt. Het gaat niet om kort "
        "plezier, maar om een duurzaam goed leven.",
    )
    add_subheading(doc, "Deugdethiek")
    add_body(
        doc,
        "Aristoteles stelt niet de vraag: 'Welke regel volg ik?', maar: 'Wat voor mens "
        "wil ik zijn?' Een deugd (aretè: goede karaktereigenschap) ontwikkel je door "
        "oefening. Het juiste handelen ligt vaak tussen twee uitersten, het midden "
        "(mesotès). Praktische wijsheid (phronesis) helpt je bepalen wat in een situatie "
        "passend is.",
    )
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "Voor professionals betekent dit: ethiek is geen checklist, maar een vaardigheid. "
        "Je ontwikkelt karakter door ervaring, reflectie en oefening. Denk bijvoorbeeld "
        "aan balans in werk en technologiegebruik — digitale gematigdheid als deugd van "
        "deze tijd.",
    )
    add_verwerkingsvraag(
        doc,
        "Aristoteles",
        "Je werkt structureel meer dan 50 uur per week en bent altijd bereikbaar. "
        "Je prestaties zijn goed, maar je voelt je uitgeput.",
        "Welke deugd ontbreekt hier? Wat is het juiste midden?",
    )

    add_image_placeholder(
        doc,
        "Aristoteles en de deugd als middenweg",
        "Visualisatie van de 'golden mean': een schaal of spectrum met uitersten "
        "(bijv. lafheid ↔ overmoed) en het midden als deugd. Optioneel: foto van "
        "Aristoteles of de Lyceum-ruïnes.",
    )

    add_subheading(doc, "Reflectievragen Deel 1")
    add_bullet(doc, "Welke aannames ('schaduwen') zie je in je vakgebied?")
    add_bullet(doc, "Welke deugd is voor jou het moeilijkst om te ontwikkelen?")

    # DEEL 2
    add_heading(doc, "Deel 2: De Middeleeuwen – De dialoog tussen rede en geloof", 1)
    add_subheading(doc, "Tijdsbeeld")
    add_body(
        doc,
        "Na de val van het West-Romeinse Rijk (476) wordt Europa grotendeels christelijk. "
        "De Kerk ontwikkelt zich tot centrum van onderwijs, zorg en macht. Tegelijkertijd "
        "blijven veel Griekse teksten bewaard in kloosters.",
    )
    add_body(
        doc,
        "Vanaf de 12e eeuw keren werken van Aristoteles via Arabische vertalingen terug in "
        "Europa. Dit leidt tot een nieuwe vraag: hoe verhouden geloof (openbaring) en rede "
        "(filosofie) zich tot elkaar?",
    )
    add_body(
        doc,
        "Voor HBO-studenten is dit een vroeg voorbeeld van interdisciplinair denken: "
        "verschillende kennisbronnen combineren binnen één vraagstuk — vergelijkbaar met "
        "hoe vandaag juridische, technische en morele kaders samenkomen (bijv. AI-regelgeving).",
    )

    add_image_placeholder(
        doc,
        "Middeleeuwse scholastiek",
        "Illustratie van een middeleeuwse scriptorium of universiteit (bijv. Parijs). "
        "Of: schema 'Geloof + Rede → synthese' met kerk en boeken als symbolen.",
    )

    # Augustinus
    add_heading(doc, "Aurelius Augustinus (354–430)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Augustinus werd geboren in Noord-Afrika (Thagaste, het huidige Algerije). In zijn "
        "jonge jaren zocht hij naar betekenis in retorica, filosofie en een werelds leven. "
        "Na een innerlijke crisis bekeerde hij zich tot het christendom en werd later bisschop.",
    )
    add_body(
        doc,
        "Hij leefde in een periode waarin het Romeinse Rijk uiteenviel. Dat riep de vraag "
        "op: waar vind je stabiliteit als maatschappelijke structuren verdwijnen?",
    )
    add_subheading(doc, "Het Goede")
    add_body(
        doc,
        "Voor Augustinus is God het hoogste goed. Aardse zaken zijn tijdelijk en kunnen "
        "geen blijvende innerlijke rust bieden. Hij beschrijft het kwaad niet als iets "
        "zelfstandigs, maar als een gebrek aan het goede — een verstoring van de juiste "
        "gerichtheid van de wil.",
    )
    add_subheading(doc, "Ethische gedachte")
    add_body(
        doc,
        "De mens neigt volgens Augustinus tot verkeerde keuzes (zonde), waarbij eigenbelang "
        "boven het goede wordt gesteld. Daarom is genade nodig: morele verbetering komt "
        "niet alleen uit eigen kracht voort.",
    )
    add_body(
        doc,
        "Deze gedachte heeft sterk bijgedragen aan het westerse idee dat intenties en "
        "innerlijke houding moreel relevant zijn.",
    )
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "In tijden van onzekerheid — bijvoorbeeld bij reorganisaties of maatschappelijke "
        "crises — stelt Augustinus een fundamentele vraag: waar baseer jij je houvast op?",
    )
    add_verwerkingsvraag(
        doc,
        "Augustinus",
        "Na een reorganisatie raken teams verdeeld en ontstaat onzekerheid. Collega's "
        "ervaren stress en verlies van richting.",
        "Waar zou hij rust zoeken: in externe zekerheid of in innerlijke oriëntatie? "
        "Hoe zou je een collega adviseren vanuit zijn denken?",
    )

    add_image_placeholder(
        doc,
        "Augustinus",
        "Portret of mozaïek van Augustinus; optioneel kaartje Noord-Afrika / Romeinse Rijk in verval.",
    )

    # Aquino
    add_heading(doc, "Thomas van Aquino (1225–1274)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Thomas van Aquino was een dominicaanse theoloog en filosoof. Hij probeerde het "
        "christelijk geloof te verbinden met de filosofie van Aristoteles. Zijn bekendste "
        "werk, de Summa Theologica, is een systematische poging om geloof en rede met "
        "elkaar te verzoenen.",
    )
    add_subheading(doc, "Het Goede")
    add_body(
        doc,
        "Volgens Aquino is het goede datgene wat overeenkomt met de natuurwet: morele "
        "principes die passen bij de menselijke natuur en via de rede te begrijpen zijn. "
        "Hij onderscheidt natuurlijke deugden (zoals rechtvaardigheid en matigheid) en "
        "theologische deugden (geloof, hoop en liefde).",
    )
    add_subheading(doc, "Ethische gedachte — natuurwet")
    add_body(
        doc,
        "De natuurwet stelt dat er universele morele principes zijn die niet afhankelijk "
        "zijn van menselijke wetten. Je ziet echo's hiervan in mensenrechten en beroepscodes.",
    )
    add_body(doc, "[CHECK: de claim van universaliteit wordt in moderne ethiek vaak betwist]")
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "In hedendaagse beroepspraktijken zie je vaak spanning tussen verschillende normen: "
        "wetgeving, professionele standaarden en persoonlijke overtuigingen. Aquino laat "
        "zien dat deze bronnen niet per se botsen, maar ook geïntegreerd kunnen worden — "
        "relevant voor ethische commissies en multidisciplinaire teams.",
    )
    add_verwerkingsvraag(
        doc,
        "Thomas van Aquino",
        "In een ziekenhuis botsen medische data, belangenorganisaties en religieuze "
        "overtuigingen bij een moreel besluit.",
        "Hoe zou Aquino rede en overtuiging laten samenwerken in één besluit? "
        "Welke deugden zijn hier in het spel?",
    )

    add_image_placeholder(
        doc,
        "Thomas van Aquino",
        "Schilderij van Aquino (vaak met zonnesymbool op borst) of diagram: "
        "natuurlijke deugden + theologische deugden.",
    )

    add_subheading(doc, "Reflectievragen Deel 2")
    add_bullet(doc, "Kun je rede en overtuiging scheiden in een professioneel dilemma?")
    add_bullet(doc, "Wat is een 'natuurwet' in jouw vakgebied — zijn er universele professionele normen?")

    # DEEL 3
    add_heading(doc, "Deel 3: De Verlichting – De autonomie van de mens", 1)
    add_subheading(doc, "Tijdsbeeld")
    add_body(
        doc,
        "In de 17e en 18e eeuw verandert Europa ingrijpend. Wetenschappelijke ontdekkingen "
        "(Newton, Galilei) en politieke ontwikkelingen versterken het vertrouwen in de "
        "menselijke rede. Autoriteit van kerk en monarchie komt onder druk te staan.",
    )
    add_body(
        doc,
        "De mens wordt gezien als een wezen dat zelf kan nadenken en keuzes kan maken. "
        "De Verlichting (Aufklärung: verlichting door rede) streeft naar vrijheid, "
        "gelijkheid en vooruitgang.",
    )

    add_image_placeholder(
        doc,
        "De Verlichting",
        "Gravure uit de 18e eeuw: salon met filosofen, of portretten van "
        "Descartes, Kant en Mill naast elkaar. Symboliek: lichtstraal, open boek, wereldbol.",
    )

    # Descartes
    add_heading(doc, "René Descartes (1596–1650)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Descartes was Frans filosoof en wiskundige. Tijdens religieuze conflicten zocht "
        "hij naar absolute zekerheid in kennis. Zijn methode was radicaal: twijfel aan "
        "alles wat betwijfeld kan worden.",
    )
    add_body(
        doc,
        "Wat uiteindelijk overblijft, is het denken zelf: cogito ergo sum — ik denk, "
        "dus ik ben.",
    )
    add_subheading(doc, "Het Goede")
    add_body(
        doc,
        "Goed handelen betekent volgens Descartes handelen op basis van heldere en "
        "zekere kennis. Passies kunnen misleiden; daarom moet de rede sturen.",
    )
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "In een tijd van data en algoritmes lijkt Descartes' nadruk op zekerheid herkenbaar. "
        "Data-gedreven besluitvorming vraagt om onderbouwing — maar wie alles wil meten, "
        "riskeert mensen als cijfers te zien. [CHECK: Kant werkt deze kritiek verder uit]",
    )
    add_verwerkingsvraag(
        doc,
        "Descartes",
        "Een AI-systeem op je stage adviseert om een klant te weigeren op basis van "
        "risicoscores. Het dashboard toont 96% betrouwbaarheid. Jouw intuïtie zegt dat "
        "het advies onterecht is, maar je kunt de berekening niet volledig uitleggen.",
        "Wat moet eerst zeker zijn voordat je handelt? Wanneer is data genoeg — en wanneer niet?",
    )

    # Kant
    add_heading(doc, "Immanuel Kant (1724–1804)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Kant leefde zijn hele leven in Königsberg (nu Kaliningrad). Hij definieerde "
        "verlichting als het uittreden uit 'zelfverschuldigde onmondigheid' — durf zelf "
        "nadenken. Zijn drie grote kritieken onderzoeken wat de rede kan weten, moet "
        "doen en mag hopen.",
    )
    add_subheading(doc, "Het Goede en de plicht")
    add_body(
        doc,
        "Volgens Kant is een handeling moreel juist wanneer deze voortkomt uit plicht, "
        "niet alleen uit resultaat. Utilitaire afweging ('het levert het meeste op') "
        "volstaat niet als je daarmee principes schendt.",
    )
    add_subheading(doc, "Categorische imperatief")
    add_body(
        doc,
        "De kernvraag: kun je jouw handelen verheffen tot algemene regel? Daarnaast: "
        "behandel mensen nooit alleen als middel, maar altijd ook als doel.",
    )
    add_body(
        doc,
        "Voor professionals: een klant, patiënt, leerling of burger is geen instrument "
        "voor jouw doel (omzet, score, efficiëntie), maar een persoon met inherente waardigheid.",
    )
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "In data-ethiek is Kant zeer relevant: gebruik van data zonder echte toestemming "
        "(informed consent) kan betekenen dat mensen gereduceerd worden tot middel. "
        "GDPR sluit hierop aan.",
    )
    add_verwerkingsvraag(
        doc,
        "Kant",
        "Je werkgever wil gebruikersdata verkopen aan adverteerders. In de algemene "
        "voorwaarden staat het 'ergens' vermeld; de meeste gebruikers hebben niet "
        "bewust ingestemd. Jij moet het technisch implementeren.",
        "Kun je dit rechtvaardigen als algemene regel? Worden mensen hier als middel "
        "behandeld? Wat zou jij doen — en waarom?",
    )

    add_image_placeholder(
        doc,
        "Kant en de categorische imperatief",
        "Portret van Kant + eenvoudig schema: 'Maxime → universaliseerbaar? → Ja/Nee'. "
        "Optioneel: icoon mens als doel (niet als middel).",
    )

    # Mill
    add_heading(doc, "John Stuart Mill (1806–1873)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Mill groeide op in het utilitarisme van zijn vader James Mill en "
        "Jeremy Bentham. In het industriële, klassengestratificeerde Engeland "
        "pleitte hij voor individuele vrijheid, vrouwenrechten en "
        "representatieve democratie. Zijn boek Over vrijheid (1859) is een "
        "klassieker van het liberale denken.",
    )
    add_subheading(doc, "Het Goede: utilitarisme")
    add_body(
        doc,
        "Het goede is het grootste geluk (nut, utility) voor het grootste aantal mensen. "
        "Morele keuzes worden beoordeeld op gevolgen, niet alleen op intenties. Mill verfijnde "
        "dit door hogere en lagere vormen van geluk te onderscheiden: culturele en "
        "intellectuele verrijking wegen zwaarder dan louter lichamelijk genot.",
    )
    add_subheading(doc, "Harmprincipe (schadeprincipe)")
    add_body(
        doc,
        "Vrijheid mag alleen worden beperkt om schade aan anderen te voorkomen. "
        "'Je mag doen wat je wilt, zolang je een ander niet direct schaadt.'",
    )
    add_body(doc, "[CHECK: 'schade' is in de praktijk ingewikkeld te definiëren]")
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "Discussies over desinformatie, online intimidatie en privacy sluiten hier direct "
        "op aan. Mill zou debat willen, maar ook ingrijpen waar vrijheid anderen schaadt.",
    )
    add_verwerkingsvraag(
        doc,
        "John Stuart Mill",
        "Een online community laat grievende reacties staan omdat 'iedereen recht heeft "
        "op een mening'. Een kleine groep leden wordt structureel getreiterd en trekt "
        "zich terug.",
        "Wanneer is ingrijpen gerechtvaardigd? Weeg het nut van vrije meningsuiting af "
        "tegen de schade aan betrokkenen.",
    )

    add_image_placeholder(
        doc,
        "Mill en het harmprincipe",
        "Illustratie: persoonlijke vrijheidssfeer vs. grens waar schade aan anderen begint. "
        "Optioneel: fabriek/industriële revolutie als contextbeeld.",
    )

    add_subheading(doc, "Reflectievragen Deel 3")
    add_bullet(doc, "Wanneer weeg jij intentie zwaarder dan resultaat — en wanneer omgekeerd?")
    add_bullet(doc, "Geef een voorbeeld uit je vakgebied waar het harmprincipe helpt bij een afweging.")

    # DEEL 4
    add_heading(doc, "Deel 4: De Moderne Tijd – Vrijheid en onzekerheid", 1)
    add_subheading(doc, "Tijdsbeeld")
    add_body(
        doc,
        "De 19e en 20e eeuw worden gekenmerkt door grote vooruitgang én grote crises: "
        "industrialisatie, kolonialisme, twee wereldoorlogen, Holocaust en atoomwapens. "
        "Vertrouwen in rede en vooruitgang komt onder druk te staan.",
    )
    add_body(
        doc,
        "God en vaste zekerheden verdwijnen voor velen uit het wereldbeeld (Nietzsche's "
        "'dood van God'). Filosofen vragen: als er geen gegeven zin is, wat betekenen "
        "vrijheid, verantwoordelijkheid en moraal dan nog?",
    )

    add_image_placeholder(
        doc,
        "Moderne tijd en existentiële crisis",
        "Collage: WO I/WO II, fabrieksrook, Holocaust-monument (respectvol), "
        "en een lege horizon — symboliseert 'verlies van zekerheid'.",
    )

    # Nietzsche
    add_heading(doc, "Friedrich Nietzsche (1844–1900)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Nietzsche was Duits filoloog en filosoof. Zijn werk werd vaak "
        "verkeerd gebruikt (o.a. door nazi-propaganda — tegen zijn bedoeling in). "
        "Hij analyseerde hoe moraal geschiedenis heeft: christelijke slavenmoraal versus "
        "aristocratische waarden. Hij vreesde dat zonder overtuigende waarden Europa in "
        "nihilisme (zinloosheid) zou vervallen.",
    )
    add_body(doc, "[CHECK: 'slavenmoraal' versus 'herenmoraal' is hier didactisch versimpeld]")
    add_subheading(doc, "Het Goede")
    add_body(
        doc,
        "Het goede is levensbevestiging, vitaliteit en creativiteit. De Übermensch "
        "(beter: 'de mens die zichzelf overstijgt') schept eigen waarden in plaats van "
        "blind overgenomen moraal te volgen.",
    )
    add_body(
        doc,
        "Nietzsche roept op tot kritiek op herdachte normen — niet tot onderdrukking "
        "van anderen. [CHECK: zijn werk werd misbruikt door latere propaganda; context is essentieel]",
    )
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "In een tijd zonder vaste autoriteiten: welke waarden kies jij zelf, en kun je "
        "die verdedigen? Vraag: leef je volgens eigen waarden of volgens overgenomen normen "
        "('best practice' zonder bevragen)?",
    )
    add_verwerkingsvraag(
        doc,
        "Nietzsche",
        "In jouw sector is 'best practice' heilig: iedereen volgt dezelfde frameworks, KPI's en "
        "certificeringen. Afwijken wordt gezien als onprofessioneel. Niemand vraagt meer waaróm "
        "deze normen goed zijn.",
        "Welke 'herdenkmoraal' zie je hier? Welke eigen waarden zou je durven formuleren "
        "zonder anderen te verachten? Wat is levensbevestigend handelen in deze context?",
    )

    # Arendt
    add_heading(doc, "Hannah Arendt (1906–1975)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Arendt ontvluchtte nazi-Duitsland en werd een invloedrijke politiek "
        "denker. Haar verslag van de proces tegen Adolf Eichmann in Jeruzalem "
        "leidde tot het begrip 'de banaliteit van het kwaad': Eichmann was "
        "geen monster, maar een bureaucraat die niet nadacht over de betekenis "
        "van zijn handelen.",
    )
    add_subheading(doc, "Het Goede")
    add_body(
        doc,
        "Het goede is politieke vrijheid: het vermogen om samen in de "
        "publieke ruimte te verschijnen, te spreken en nieuw te beginnen "
        "(natality). Moraal is niet alleen privé, maar verschijnt in "
        "gezamenlijk handelen en oordelen.",
    )
    add_subheading(doc, "Banaliteit van het kwaad")
    add_body(
        doc,
        "Kwaad kan ontstaan zonder slechte intentie — door gedachteloos handelen, routines "
        "volgen en verantwoordelijkheid afschuiven ('ik volgde orders'). Het goede vereist "
        "kritisch burgerschap en oplettendheid, ook in organisaties.",
    )
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "Algoritmes die ons in bubbels plaatsen, vernietigen de publieke ruimte. "
        "In organisaties blijft persoonlijke verantwoordelijkheid bestaan: 'Ik deed alleen "
        "mijn taak' is geen excu als je had kunnen weten wat er misging.",
    )
    add_verwerkingsvraag(
        doc,
        "Hannah Arendt",
        "In een grote organisatie volgt iedereen een nieuw IT-protocol dat tijd bespaart. "
        "Jij ziet dat het kwetsbare cliënten systematisch overslaat, maar collega's zeggen: "
        "'Ik voer alleen uit wat het systeem vraagt.'",
        "Waar is hier sprake van 'banaliteit van het kwaad'? Hoe herstel je denken en "
        "verantwoordelijkheid in de publieke/professionele ruimte? Wat zou jij concreet doen?",
    )

    add_image_placeholder(
        doc,
        "Hannah Arendt en de publieke ruimte",
        "Portret Arendt + schema 'publieke ruimte' (agora/marktplein metafoor) "
        "versus afgesloten bubbels. Respectvol beeld bij Eichmann-proces (geen sensatie).",
    )

    # Sartre & Beauvoir
    add_heading(doc, "Jean-Paul Sartre (1905–1980) & Simone de Beauvoir (1908–1986)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Sartre en Beauvoir waren partners en leidende existentialisten in "
        "naoorlogs Parijs. De Tweede Wereldoorlog toonde dat mensen radicaal "
        "vrij zijn — ook om te kiezen voor medeplichtigheid of verzet. "
        "Beauvoir werd feministe avant la lettre met Le Deuxième Sexe (1949).",
    )
    add_subheading(doc, "Het Goede")
    add_body(
        doc,
        "De mens is 'veroordeeld tot vrijheid' (Sartre): je kunt niet niet kiezen. "
        "Goed leven is authentiek leven — niet wegkruipen achter rollen, gewoonten of "
        "excuses ('de ander maakt mij zo').",
    )
    add_body(
        doc,
        "Beauvoir voegde toe dat veel 'natuurlijke' verschillen (zoals vrouwelijkheid) "
        "sociaal geconstrueerd zijn: 'men wordt geen vrouw, men wordt het'.",
    )
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "In je carrière: welke keuzes maak jij echt zelf, en waar verstop "
        "je je achter 'zo hoort het'? Beauvoir helpt gender en macht in "
        "beroepscontexten te analyseren.",
    )
    add_verwerkingsvraag(
        doc,
        "Sartre en Beauvoir",
        "Op je stage hoor je regelmatig sexistische grappen; de leiding zegt: 'Zo is de cultuur hier, "
        "ga je niet aanpassen, ga je maar weg.' Je hebt de stage nodig voor je diploma.",
        "Waar verstop je je achter mauvaise foi (slechte trouw / zelfbedrog) als je zwijgt? "
        "Wat zou Beauvoir zeggen over sociaal geconstrueerde 'normaliteit'? "
        "Formuleer een authentieke keuze — mét de consequenties die je accepteert.",
    )

    # Rawls
    add_heading(doc, "John Rawls (1921–2002)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Rawls was Amerikaans politiek filosoof. Na de crisis van "
        "rechtvaardigheid in de jaren '60 zocht hij een rationeel "
        "fundament voor een eerlijke samenleving — niet gebaseerd op "
        "particuliere moraal of utilitaire rekenkunde alleen.",
    )
    add_subheading(doc, "Het Goede: rechtvaardigheid als billijkheid")
    add_body(
        doc,
        "Rechtvaardigheid ontstaat wanneer je regels bedenkt alsof je niet weet welke "
        "positie je zelf in de samenleving krijgt (het sluier van onwetendheid). "
        "Achter dit sluier kies je principes die de vrijste basismogelijkheden voor "
        "iedereen waarborgen.",
    )
    add_body(
        doc,
        "Ongelijkheid is alleen toegestaan als die de minstbedeelden ten goede komt "
        "(het verschilbeginsel).",
    )
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "Beleid ontwerpen alsof je zelf de zwakste positie kunt krijgen: "
        "toegankelijkheid, inclusie, eerlijke kansen in onderwijs en "
        "arbeidsmarkt. Relevant voor HR-beleid en publieke dienstverlening.",
    )
    add_verwerkingsvraag(
        doc,
        "John Rawls",
        "Je team mag een nieuw stagebeleid ontwerpen: wie krijgt de beste plekken, wie betaalt "
        "het reisgeld, wie mag thuiswerken? Jij weet dat jij volgend semester teamleider wordt.",
        "Ontwerp het beleid alsof je achter Rawls' sluier van onwetendheid staat: je weet niet "
        "welke positie je krijgt. Welke regels kies je en waarom zijn die billijk?",
    )

    add_image_placeholder(
        doc,
        "Rawls en het sluier van onwetendheid",
        "Illustratie: figuren achter een sluier/ gordijn, symbolisch gelijke "
        "onwetendheid over status. Diagram: basisstructuur van samenleving + "
        "verschilbeginsel in eenvoudige vorm.",
    )

    # Camus
    add_heading(doc, "Albert Camus (1913–1960)", 2)
    add_subheading(doc, "Achtergrond")
    add_body(
        doc,
        "Camus groeide op in armoede in Algerije en werd schrijver en "
        "filosoof in Frankrijk. In De mythe van Sisyphus onderzoekt hij "
        "het absurde: de mens verlangt naar zin, het universum antwoordt "
        "niet. Zijn antwoord is geen zelfmoord of wanhoop, maar rebellie "
        "en mede-menselijkheid.",
    )
    add_subheading(doc, "Het Goede")
    add_body(
        doc,
        "Het goede is solidariteit en verzet tegen onrecht, ook als het universum "
        "onverschillig blijft. Het leven is niet vanzelf zinvol — maar krijgt betekenis "
        "door handelen.",
    )
    add_body(
        doc,
        "Sisyphus — veroordeeld om eeuwig een rots de berg op te duwen — moet je je als "
        "gelukkig mens voorstellen: hij erkent het absurde en blijft toch handelen.",
    )
    add_subheading(doc, "Moderne vertaalslag")
    add_body(
        doc,
        "In zorg, onderwijs of maatschappelijk werk: soms lijkt je inzet "
        "een druppel op een gloeiende plaat. Camus zegt: de strijd zelf "
        "heeft waarde — solidariteit blijft nodig.",
    )
    add_verwerkingsvraag(
        doc,
        "Albert Camus",
        "Je werkt in de jeugdzorg. Wachtlijsten groeien, systemen falen, je ziet dezelfde patronen "
        "terug bij elke casus. Collega's zeggen: 'Het heeft toch geen zin.'",
        "Waarom zou Camus zeggen dat je tóch moet blijven handelen? Wat is 'Sisyphus als "
        "gelukkig mens' in jouw beroep? Beschrijf één daad van solidariteit die zin heeft "
        "óók zonder gegarandeerd succes.",
    )

    add_subheading(doc, "Reflectievragen Deel 4")
    add_bullet(doc, "Waar zie je 'banale' morele blindheid in organisaties?")
    add_bullet(doc, "Welke samenleving zou je ontwerpen achter Rawls' sluier?")

    # DEEL 5
    add_heading(doc, "Deel 5: Reflectie op de digitale agora", 1)
    add_body(
        doc,
        "Nieuwe technologie vormt een moderne 'agora' (marktplein): een plek waar informatie, "
        "meningen en macht samenkomen. Oude filosofische vragen keren terug in nieuwe vorm: "
        "wat is waarheid? wat is verantwoordelijkheid? wat is vrijheid?",
    )
    add_body(
        doc,
        "Voor HBO-professionals in technologie, zorg, onderwijs en communicatie bieden "
        "de denkers uit dit manuscript scherpe 'brilwijzers' voor actuele dilemma's.",
    )

    add_image_placeholder(
        doc,
        "De digitale agora",
        "Moderne illustratie: smartphones en laptops vormen een 'marktplein'; "
        "datastromen als schaduwen op een muur (referentie Plato). "
        "Subtiel, niet technofobisch.",
    )

    add_subheading(doc, "Vijf brilwijzers voor vandaag")
    add_bullet(
        doc,
        "Arendt: algoritmes die ons in bubbels plaatsen, vernietigen de publieke ruimte — "
        "democratie vereist gedeelde werkelijkheid en gesprek.",
    )
    add_bullet(
        doc,
        "Aristoteles: digitale gematigdheid als deugd van deze eeuw — balans tussen "
        "online en offline, tussen bereikbaarheid en rust.",
    )
    add_bullet(
        doc,
        "Kant: data voor winst zonder transparantie behandelt mensen als middel — "
        "informed consent en menselijke waardigheid zijn niet optioneel.",
    )
    add_bullet(
        doc,
        "Mill: desinformatie en schade — waar ligt de grens tussen vrije meningsuiting "
        "en ingrijpen? Vraag naar aantoonbare schade en proportionele maatregelen.",
    )
    add_bullet(
        doc,
        "Socrates: stel vragen aan AI-output en aan je eigen aannames — wie controleert "
        "de kwaliteit van wat je overneemt?",
    )

    add_subheading(doc, "Casus voor groepswerk (HBO)")
    add_body(
        doc,
        "Een zorginstelling wil AI inzetten voor risicoscreening van patiënten. "
        "Behandelteams twijfelen over privacy, bias in data en verantwoordelijkheid "
        "bij fouten. Gebruik minimaal drie denkers uit dit document om argumenten "
        "voor en tegen te formuleren. Presenteer een gemotiveerd advies.",
    )

    add_image_placeholder(
        doc,
        "Casus AI in de zorg",
        "Infographic: stakeholderkaart (patiënt, zorgverlener, IT-leverancier, overheid) "
        "rond een AI-systeem. Geen specifiek merk; neutraal en didactisch.",
    )

    # VERGELIJKEND OVERZICHT
    add_heading(doc, "Overzicht: Hoe denken deze filosofen over 'het goede'?", 1)
    add_body(
        doc,
        "Socrates — kennis en inzicht | Plato — het Goede als hoogste Idee | "
        "Aristoteles — eudaimonia en deugd | Augustinus — liefde voor God | "
        "Aquino — natuurwet en deugden | Descartes — rede stuurt passies | "
        "Kant — plicht en menselijke waardigheid | Mill — grootste geluk, minste schade | "
        "Nietzsche — levenskracht en eigen waarden | Arendt — politieke vrijheid | "
        "Sartre/Beauvoir — authenticiteit | Rawls — billijke verdeling | "
        "Camus — solidariteit en rebellie",
    )

    add_image_placeholder(
        doc,
        "Vergelijkende matrix",
        "Tabel of matrix: denkers (rijen) × thema's (intentie, gevolg, deugd, plicht, vrijheid). "
        "Kleurcodering per periode. Geschikt als poster in de klas.",
    )

    # AANVULLING STROMINGEN
    add_heading(doc, "Aanvulling: Stromingen die het huidige tijdsgewricht helpen verklaren", 1)
    add_body(
        doc,
        "Het document behandelt grote denkers en hun ethiek, maar niet elke belangrijke "
        "stroming. Onderstaande bewegingen ontbreken in de hoofdlijn — toch zijn ze "
        "zeer nuttig om het heden te begrijpen. Per stroming: waarom het ertoe doet, "
        "kernbegrip, en link met je beroepspraktijk.",
    )

    add_image_placeholder(
        doc,
        "Kaart van aanvullende stromingen",
        "Mindmap of clusterdiagram: centrum 'Heden' met takken naar pragmatisme, "
        "fenomenologie, kritische theorie, poststructuralisme, zorgethiek, communitarisme, "
        "capabilities, decoloniaal denken, techniekfilosofie, klimaatethiek.",
    )

    stromingen = [
        (
            "Pragmatisme (Peirce, James, Dewey)",
            "Filosofie die vraagt: wat werkt in de praktijk? Waarheid is niet alleen "
            "abstract, maar iets dat zich bewijst in ervaring en handelen. Dewey's "
            "idee van 'leren door doen' is direct herkenbaar in HBO-onderwijs en "
            "reflectief stagelopen.",
            "Nut voor nu: ethiek is geen theorieboek alleen — test je morele keuzes "
            "in de praktijk en pas ze aan op basis van wat je leert.",
            "Casus: een nieuw beleid ziet er op papier goed uit, maar werkt niet in de "
            "praktijk. Pragmatisme vraagt: welke consequenties meet je, en hoe verbeter je?",
        ),
        (
            "Fenomenologie (Husserl, Merleau-Ponty)",
            "Onderzoekt hoe de wereld verschijnt in levendige ervaring — niet vanuit "
            "buitenaf, maar vanuit de eerste persoon. Lichaam, zorg, perceptie en "
            "context tellen mee.",
            "Nut voor nu: in zorg, onderwijs en design: wat is de ervaring van de "
            "patiënt, leerling of gebruiker — niet alleen wat het protocol zegt?",
            "Casus: een app is technisch perfect, maar gebruikers voelen zich niet "
            "gehoord. Fenomenologie vraagt: hoe ervaren zij het systeem?",
        ),
        (
            "Kritische theorie / Frankfurter Schule (Horkheimer, Adorno, Habermas)",
            "Analyseert hoe macht, kapitaal en cultuur ons denken vormen. Adorno "
            "waarschuwde voor de 'cultuurindustrie'; Habermas voor de 'instrumentele "
            "reden' die alles tot middel maakt. Communicatief handelen vraagt om "
            "eerlijk gesprek zonder manipulatie.",
            "Nut voor nu: platforms, algoritmes en marketing zijn geen neutrale tools — "
            "ze sturen gedrag. Kritische theorie helpt vragen: wiens belang dient dit systeem?",
            "Casus: een bedrijf presenteert AI als 'objectief', maar verdient aan "
            "meer gebruik. Wie profiteert — en wie betaalt de prijs?",
        ),
        (
            "Poststructuralisme (Foucault, Lyotard)",
            "Foucault: kennis en macht horen bij elkaar — instituten bepalen wat "
            "'normaal' en 'waar' heet. Lyotard: wantrouwen tegen 'grote verhalen' "
            "die alles verklaren (vooruitgang, wetenschap, vrijemarkt).",
            "Nut voor nu: begrijp waarom sommige stemmen gehoord worden en andere niet. "
            "Handig bij debatten over diversiteit, surveillance en institutioneel racisme.",
            "Casus: een organisatie zegt 'we zijn inclusief', maar promotiepatronen "
            "vertellen een ander verhaal. Foucault vraagt: welke machtsstructuren "
            "produceren dit 'normale'?",
        ),
        (
            "Zorgethiek / ethics of care (Gilligan, Tronto, Held)",
            "Tegenhanger van abstracte regel-ethiek: moraal ontstaat in relaties van "
            "afhankelijkheid en zorg. Tronto onderscheidt fasen: zorgzaamheid opmerken, "
            "verantwoordelijkheid nemen, zorg verlenen, responsiviteit.",
            "Nut voor nu: essentieel in zorg, opvoeding, HR en maatschappelijk werk — "
            "waar mensen kwetsbaar zijn en niet reduceerbaar tot dossiers.",
            "Casus: een cliënt wil iets dat 'niet in het protocol past'. Zorgethiek "
            "vraagt: welke relatie en welke zorg vraagt deze persoon nu?",
        ),
        (
            "Communitarisme (MacIntyre, Taylor, Sandel)",
            "Kritiek op het idee van de volledig autonome individuele consument. "
            "Menselijke identiteit en moraal groeien in gemeenschappen, tradities "
            "en gedeelde praktijken ('practices'). MacIntyre herneemt deugdethiek: "
            "deugd is verbonden met vakmanschap in een leefbare gemeenschap.",
            "Nut voor nu: sluit aan bij 'moreel vakmanschap' en beroepsidentiteit — "
            "je bent niet alleen individu, maar lid van een praktijk met normen.",
            "Casus: een professional moet kiezen tussen bedrijfsbelang en "
            "beroepsethiek. Communitarisme vraagt: welke gemeenschap en welke "
            "traditie geven je richting?",
        ),
        (
            "Capability approach (Sen, Nussbaum)",
            "Rechtvaardigheid gaat niet alleen over geld of formele rechten, maar "
            "over reële mogelijkheden om te leven zoals je waardeert (capabilities). "
            "Nussbaum formuleerde een lijst van centrale menselijke vermogens.",
            "Nut voor nu: beleid en design toetsen op wat mensen echt kúnnen doen — "
            "niet alleen wat ze op papier mogen.",
            "Casus: een stad is 'toegankelijk' op papier, maar slecht bereikbaar "
            "voor mensen met weinig geld of tijd. Wat zijn hun reële kansen?",
        ),
        (
            "Decoloniaal denken (Fanon, Maldonado-Torres, Mbembe)",
            "Onderzoekt hoe koloniale machtsrelaties doorwerken in kennis, instituten "
            "en identiteit. Niet alleen 'geschiedenis', maar structuren die nu nog "
            "bepalen wie gezag heeft en welke kennis 'objectief' heet.",
            "Nut voor nu: essentieel voor diversiteit, globalisering, technologie-export "
            "en kritisch kijken naar 'universele' normen die Westerse oorsprong hebben.",
            "Casus: een AI-model is getraind op voornamelijk Engelstalige data en "
            "presteert slechter op andere talen en culturen. Wiens wereldbeeld is "
            "'normaal'?",
        ),
        (
            "Techniekfilosofie (Heidegger, Jonas, Floridi)",
            "Heidegger: technologie is geen neutraal gereedschap maar een manier "
            "waarop de wereld voor ons verschijnt. Jonas: nieuwe techniek vereist "
            "een 'ethiek van verantwoordelijkheid' voor toekomstige generaties. "
            "Floridi: informatie-ethiek voor het digitale tijdperk.",
            "Nut voor nu: AI, biotech en klimaattechnologie dwingen tot vooruitdenken: "
            "wat doet deze techniek met ons — niet alleen wat wij ermee doen?",
            "Casus: een ziekenhuis introduceert AI-diagnostiek. Techniekfilosofie "
            "vraagt: verandert dit wat 'zorg' en 'diagnose' voor ons betekenen?",
        ),
        (
            "Klimaat- en generatie-ethiek (Parfit, Jonas, jongere denkers)",
            "Morele vragen over verantwoordelijkheid voor toekomstige generaties, "
            "ecologische grenzen en collectieve schade door individueel gedrag.",
            "Nut voor nu: elk beroep raakt aan duurzaamheid — van inkoop tot IT "
            "(energieverbruik van data centers) tot zorg (hitte, luchtkwaliteit).",
            "Casus: je bedrijf wil goedkoopste leverancier kiezen met hoge CO₂-uitstoot. "
            "Welke verantwoordelijkheid draag jij richting mensen die de gevolgen "
            "nu nog niet zien?",
        ),
    ]

    for naam, uitleg, nut, casus in stromingen:
        add_heading(doc, naam, 2)
        add_subheading(doc, "Kern")
        add_body(doc, uitleg)
        add_subheading(doc, "Waarom nuttig voor het heden")
        add_body(doc, nut)
        add_subheading(doc, "Mini-casus")
        add_body(doc, casus)

    add_subheading(doc, "Reflectie: welke stroming past bij jouw vakgebied?")
    add_body(
        doc,
        "Kies twee stromingen uit dit aanvullende hoofdstuk en leg uit welke vragen "
        "zij openen in jouw studie of stage — en welke denker uit het hoofddeel "
        "daarmee in gesprek gaat (bijv. zorgethiek + Aristoteles; kritische theorie + Arendt).",
    )

    # BRONNEN
    add_heading(doc, "Bronvermelding (APA 7)", 1)
    refs = [
        "Arendt, H. (2019). De menselijke conditie. (C. Houwaard, Vert.). Uitgeverij Boom.",
        "Bakewell, S. (2016). De existentialisten: Over vrijheid, zijn en cocktails. Uitgeverij Ten Have.",
        "Camus, A. (2023). De mythe van Sisyphus. (H. de Man, Vert.). IJzer.",
        "Kant, I. (2009). Kritiek van de praktische rede. (J. Veenbaas & W. Visser, Vert.). Boom.",
        "Mill, J. S. (2012). Over vrijheid. (W. E. Krul, Vert.). Boom.",
        "Rawls, J. (2006). Een theorie van rechtvaardigheid. (F. Bestebreurtje, Vert.). Lemniscaat.",
        "Störig, H. J. (2017). Geschiedenis van de filosofie. (T. van der Slikke, Vert.; 30ste druk). Uitgeverij Unieboek | Het Spectrum.",
        "Foucault, M. (2018). Toezicht en straf. (R. van der Veer, Vert.). Boom.",
        "Held, V. (2006). The ethics of care: Personal, political, and global. Oxford University Press.",
        "Nussbaum, M. C. (2011). Creating capabilities: The human development approach. Harvard University Press.",
        "Sandel, M. J. (2010). Rechtvaardigheid: Een politieke filosofie. (H. Driessen, Vert.). Uitgeverij Ten Have.",
        "Tronto, J. C. (2013). Caring democracy: Markets, equality, and justice. NYU Press.",
    ]
    for r in refs:
        add_bullet(doc, r)

    add_body(
        doc,
        "Aanvullende bronnen voor verdieping (HBO): Plato, Politeia (vertaling Boom of "
        "Prometheus); Aristoteles, Nicomacheaanse ethiek; Beauvoir, Het tweede geslacht; "
        "Nietzsche, Zo sprak Zarathustra (selectief en met begeleiding). "
        "Online: Stanford Encyclopedia of Philosophy (plato.stanford.edu) — Engelstalig maar betrouwbaar.",
    )

    doc.save(OUT)
    print(f"Opgeslagen: {OUT}")


if __name__ == "__main__":
    main()

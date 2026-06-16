# -*- coding: utf-8 -*-
"""Genereer redactie-prompt document voor filosofie/ethiek manuscript."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = (
    r"C:\Users\876409\OneDrive - Office 365 Fontys\Documenten"
    r"\Ethiek en filpsofie lessenreeks\AI_Redactie_prompt_filosofie_en_ethiek.docx"
)


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def body(doc, text):
    doc.add_paragraph(text)


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def code_block(doc, text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("AI-redactie prompts", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body(
        doc,
        "Kant-en-klare prompts voor Claude, ChatGPT of Copilot — "
        "bij het manuscript Een geschiedenis van de filosofie en ethiek (HBO).",
    )

    heading(doc, "Hoe gebruik je dit document?", 1)
    bullet(doc, "Maak een Project/Custom GPT en upload je manuscript (.docx of .pdf).")
    bullet(doc, "Plak eerst de HOOFD-PROMPT (sectie 1) als vaste instructie.")
    bullet(doc, "Gebruik daarna een TAAK-PROMPT (sectie 2–6) per klus.")
    bullet(doc, "Controleer altijd feiten, citaten en bronnen zelf — AI hallucineert.")
    bullet(doc, "Markeer AI-voorstellen als concept tot jij ze goedkeurt.")

    heading(doc, "1. Hoofd-prompt (vaste instructie)", 1)
    body(doc, "Kopieer dit naar ‘Instructions’ / ‘Custom instructions’ / systeemprompt:")
    code_block(
        doc,
        """JE ROL
Je bent senior redacteur en didactisch ontwerper voor HBO-onderwijs (Fontys-niveau).
Je werkt aan het manuscript "Een geschiedenis van de filosofie en ethiek" — een overzicht
van westerse filosofie en ethiek voor HBO-studenten, met:
- historische context per periode
- denkerprofielen (achtergrond, het goede, ethische gedachte, moderne vertaalslag)
- verwerkingsvragen met casussen per denker
- een hoofdstuk over aanvullende stromingen (pragmatisme, fenomenologie, kritische theorie, enz.)
- reflectie op de digitale agora (AI, algoritmes, sociale media)

DOELGROEP
- HBO-studenten (18–25), diverse studies (zorg, IT, onderwijs, business, communicatie)
- Taalniveau: B1/B2 Nederlands — helder, niet kinderachtig, niet wo-promoveren
- Lezers hebben géén filosofische voorkennis

SCHRIJFSTIJL
- Zakelijk-warm, uitnodigend tot nadenken
- Korte alinea's (max. 4–5 zinnen)
- Actieve zinnen, concrete voorbeelden uit studie, stage en beroepspraktijk
- Vermijd jargon; leg vaktermen bij eerste gebruik uit in 1 zin
- Geen opsommingen zonder inleidende zin
- Geen clichés ("sinds de oertijd", "in deze snelle wereld")
- Geen moraliserende toon ("je moet gewoon…")

INHOUDELIJKE REGELS
- Wijk niet af van de canon in het manuscript tenzij ik het expliciet vraag
- Voeg GEEN nieuwe denkers, citaten of bronnen toe zonder ze te markeren met [BRON NODIG]
- Parafraseer filosofen; citeer alleen met exacte bron als die in het document staat
- Houd westerse canon + kritische reflectie in balans (inclusief zorgethiek, decoloniaal, enz.)
- Moderne vertaalslagen moeten herkenbaar zijn voor HBO: stage, team, platform, AI, zorg
- Respectvol bij zwaar materiaal (Holocaust, Eichmann, kolonialisme)

OUTPUT
- Schrijf in het Nederlands
- Geef aan wat je hebt gewijzigd (kort, in bullets) als je redigeert
- Bij twijfel over feiten: schrijf [CHECK: …] in plaats van te gokken
- Behoud bestaande structuur en koppen tenzij ik anders vraag

WAT JE NIET DOET
- Geen feitelijke filosofische claims verzinnen
- Geen APA-bronnen fabriceren
- Geen tekst verkorten tot staccato of opblazen tot academisch proza
- Geen politieke of religieuze stellingname buiten wat de denker zelf zou kunnen dragen""",
    )

    heading(doc, "2. Taak-prompt: algemene redactie", 1)
    body(doc, "Plak onder de hoofd-prompt + selecteer een hoofdstuk of plak tekst:")
    code_block(
        doc,
        """Redigeer de onderstaande tekst voor HBO-niveau.

Taken:
1. Maak zinnen helderder en korter waar nodig
2. Verwijder herhaling en overbodige bijwoorden
3. Behoud alle feitelijke inhoud en structuur
4. Zorg dat vaktermen kort worden uitgelegd
5. Markeer twijfelpunten met [CHECK: …]

Geef daarna in max. 5 bullets aan wat je belangrijkste wijzigingen waren.

TEKST:
[plak hier]""",
    )

    heading(doc, "3. Taak-prompt: didactisch aanvullen", 1)
    code_block(
        doc,
        """Vul het onderstaande hoofdstuk/denkerprofiel didactisch aan voor HBO-studenten.

Voeg toe (waar passend, niet overladen):
- 1 korte "Waarom dit nu relevant is"-alinea
- 1 verdiepingskader (bijv. "In de praktijk betekent dit…")
- 1 common misconception + correctie ("Veel studenten denken…, maar…")
- Suggestie voor [AFBEELDING: …] als die ontbreekt

Maximaal 300 woorden extra. Geen nieuwe denkers zonder [BRON NODIG].

TEKST:
[plak hier]""",
    )

    heading(doc, "4. Taak-prompt: verwerkingsvraag + casus", 1)
    code_block(
        doc,
        """Maak 1 nieuwe verwerkingsvraag voor denker: [NAAM DENKER]

Formaat (exact aanhouden):
**Verwerkingsvraag — vanuit [denker]**
**Casus:** [herkenbare HBO-situatie: stage, team, social media, AI, zorg, HR — 3–5 zinnen]
**Wat zou je vanuit het denken van [denker] nu doen?** [1–2 zinnen instructie + 2 deelvragen]

Eisen:
- Casus moet moreel spannend zijn (geen zwart-wit)
- Sluit aan bij kernidee van de denker (niet generiek "wat is ethiek?")
- Geschikt voor schriftelijke opdracht (½ A4) of paargesprek (15 min)
- Geen herhaling van bestaande casussen in het manuscript

Bestaande casus ter vermijding van overlap:
[kort samenvatten of "zie manuscript"]""",
    )

    heading(doc, "5. Taak-prompt: moderne vertaalslag scherpen", 1)
    code_block(
        doc,
        """Herschrijf de "Moderne vertaalslag" voor [NAAM DENKER].

Koppel aan minstens 2 concrete contexten:
- digitale technologie (AI, algoritmes, data)
- beroepspraktijk HBO (stage, team, cliënt, klant, collega)

Eisen:
- Max. 120 woorden
- Geen simplistische "X zou TikTok haten"-claims
- Toon hoe de denker een BRIL is, geen voorspeller van apps

Huidige tekst:
[plak hier]""",
    )

    heading(doc, "6. Taak-prompt: bronnen en feitencheck", 1)
    code_block(
        doc,
        """Voer een feitencheck uit op onderstaande tekst.

Per claim:
- ✅ Waarschijnlijk correct (met korte toelichting)
- ⚠️ Nuance nodig (leg uit)
- ❌ Waarschijnlijk onjuist of oversimplificatie (stel correctie voor)
- [BRON NODIG] als je het niet kunt verifiëren

Controleer specifiek:
- jaartallen en biografische feiten
- parafrases van kernideeën (Socrates, Kant, Foucault, enz.)
- of moderne vertaalslagen de denker niet vervalsen

Geen bronnen verzinnen. Bij twijfel: [BRON NODIG].

TEKST:
[plak hier]""",
    )

    heading(doc, "7. Taak-prompt: samenwerking tussen denkers", 1)
    body(doc, "Handig voor groepsopdrachten of verdieping:")
    code_block(
        doc,
        """Ontwerp een groepsopdracht waarin twee denkers "in gesprek" gaan over deze casus:

CASUS: [bijv. AI-risicoscreening in de zorg / dataverkoop / filterbubbels / stagecultuur]

Denkers: [DENKER A] en [DENKER B]

Lever:
1. Korte casusbeschrijving (max. 150 woorden)
2. Standpunt denker A (max. 100 woorden, in hun logica)
3. Standpunt denker B (max. 100 woorden)
4. 3 discussievragen voor studenten
5. Beoordelingshint voor docent (waar op letten in antwoorden)

Geen karikaturen: beide denkers moeten sterk klinken.""",
    )

    heading(doc, "8. Taak-prompt: afbeeldingssuggestie", 1)
    code_block(
        doc,
        """Voor hoofdstuk/denker: [NAAM]

Geef 1 afbeeldingssuggestie in dit formaat:
- Titel: [AFBEELDING: …]
- Beschrijving (voor ontwerper): …
- Zoektermen (Wikimedia/Unsplash): …
- Adobe Firefly-prompt (Engels): …
- Tip: Firefly of echt beeld?

Sluit aan bij stijl: educatief, rustig, HBO-waardig, geen cliché denkbeeld.""",
    )

    heading(doc, "9. Workflow: aanbevolen volgorde", 1)
    steps = [
        "Upload manuscript in Claude Project of ChatGPT Project.",
        "Plak hoofd-prompt als vaste instructie.",
        "Feitencheck per hoofdstuk (prompt 6).",
        "Algemene redactie per hoofdstuk (prompt 2).",
        "Didactisch aanvullen waar dun (prompt 3).",
        "Extra verwerkingsvragen waar gewenst (prompt 4).",
        "Jij: eindredactie, bronnen, goedkeuring.",
        "Afbeeldingen: Firefly + Afbeeldingenlijst_filosofie_en_ethiek.docx.",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}", style="List Number")

    heading(doc, "10. Kwaliteitschecklist (mens — altijd zelf doen)", 1)
    checks = [
        "Klopt de parafrase van de denker (niet te modern, niet te vereenvoudigd)?",
        "Is het HBO-niveau (niet te academisch, niet te simpel)?",
        "Zijn casussen herkenbaar en respectvol?",
        "Zijn citaten en bronnen echt en correct?",
        "Past de toon bij Fontys / moreel vakmanschap?",
        "Is AI-gebruik gedocumenteerd volgens instellingsbeleid?",
    ]
    for c in checks:
        bullet(doc, c)

    heading(doc, "11. Voorbeeld: korte opdracht om te testen", 1)
    body(doc, "Test of je prompt werkt met deze mini-opdracht:")
    code_block(
        doc,
        """Redigeer alleen de inleiding (eerste 2 alinea's) van het manuscript.
Houd de toon uitnodigend. Max. 10% korter. Geef daarna 3 bullets met je wijzigingen.""",
    )

    doc.save(OUT)
    print(f"Opgeslagen: {OUT}")


if __name__ == "__main__":
    main()

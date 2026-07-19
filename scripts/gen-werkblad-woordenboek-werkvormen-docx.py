#!/usr/bin/env python3
"""Genereer MV_21 — Werkvormen bij het Moreel Woordenboek (NL + EN)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "public/downloads/zien"

C_TITLE = RGBColor(0x1A, 0x27, 0x44)
C_ACCENT = RGBColor(0x18, 0x5F, 0xA5)
C_BODY = RGBColor(0x5F, 0x5E, 0x5A)

CONTENT = {
    "nl": {
        "filename": "MV_21_MoreelWoordenboek_Werkvormen_NL.docx",
        "title": "Werkvormen bij het Moreel Woordenboek",
        "subtitle": "MV_21  ·  Stap Zien  ·  Community Moreel Vakmanschap",
        "intro": (
            "Het Moreel Woordenboek verzamelt begrippen die vaak voorkomen in beleid, "
            "gedragscodes en gesprekken over integriteit en sociale veiligheid. "
            "Deze vier werkvormen helpen om die taal zichtbaar, bespreekbaar en toetsbaar te maken — "
            "ook in het licht van espoused theory versus theory-in-use (Argyris & Schön)."
        ),
        "forms": [
            {
                "title": "1. Woordenjacht in beleid",
                "meta": "20–30 min  ·  individueel of duo’s  ·  Zien",
                "doel": "Zichtbaar maken welke morele woorden een beleids- of gedragstekst dominant gebruikt.",
                "stappen": [
                    "Kies een korte passage uit een gedragscode, integriteitsbeleid of teamafspraak (½–1 pagina).",
                    "Markeer alle woorden die in het Moreel Woordenboek voorkomen — of die moreel ‘aanvoelen’, ook als ze er nog niet in staan.",
                    "Noteer de top 5 meest voorkomende of meest beladen termen.",
                    "Bespreek: welk moreel frame spreekt deze tekst? Regels, zorg, moed, veiligheid, inclusie…?",
                ],
                "reflectie": [
                    "Welke woorden mist je die wél bij jullie praktijk horen?",
                    "Welke term zou je schrappen als je de tekst eerlijker zou maken?",
                ],
            },
            {
                "title": "2. Espoused vs. in use",
                "meta": "30–40 min  ·  3–6 personen  ·  Zien → Voelen",
                "doel": "Het verschil onderzoeken tussen wat een tekst claimt (espoused) en wat praktijk laat zien (theory-in-use).",
                "stappen": [
                    "Iedereen kiest uit het woordenboek 3 termen die ‘op papier’ bij de eigen organisatie of opleiding horen.",
                    "Schrijf per term één concreet voorbeeld waarin die waarde wél zichtbaar is, en één waarin die juist ontbreekt of wringt.",
                    "Deel in de groep alleen de wringende voorbeelden — zonder herleidbare details.",
                    "Formuleer samen één zin: ‘Wij zeggen X, maar in de praktijk belonen we vaak Y.’",
                ],
                "reflectie": [
                    "Welke term bleek espoused, maar zelden in use?",
                    "Wat zou één kleine gedragsverandering zijn die de kloof verkleint?",
                ],
            },
            {
                "title": "3. Clusterdebat",
                "meta": "25–35 min  ·  4–8 personen  ·  Wegen",
                "doel": "Ontdekken welk moreel cluster in een situatie of tekst het hardst trekt — en wat daardoor buiten beeld raakt.",
                "stappen": [
                    "De begeleider noemt een herkenbare situatie of leest een korte beleidspassage voor.",
                    "De groep verdeelt zich over clusters: integriteit/compliance, zorg/relatie, moed/verantwoordelijkheid, veiligheid, diversiteit/inclusie, transparantie/vertrouwen.",
                    "Elk cluster beargumenteert in 2 minuten waarom ‘hun’ frame hier dominant zou moeten zijn.",
                    "Daarna: welk cluster was ondervertegenwoordigd, en wat kost dat moreel?",
                ],
                "reflectie": [
                    "Welke cluster wint meestal in jullie documenten?",
                    "Welke cluster wint in jullie dagelijkse besluiten?",
                ],
            },
            {
                "title": "4. Lexicon-uitbreiding",
                "meta": "15–20 min  ·  plenair of kleine groepen  ·  Zien",
                "doel": "Het woordenboek levend houden: ontbrekende termen aandragen die in beleid of praktijk ertoe doen.",
                "stappen": [
                    "Iedereen noteert 1–2 morele woorden die wél in jullie praktijk spelen, maar (nog) niet in het lexicon staan.",
                    "Per term: korte werkdefinitie in één zin + waarom die term nodig is.",
                    "De groep stemt: opnemen, later bespreken, of afwijzen.",
                    "Goedgekeurde termen gaan naar de redactie van het Moreel Woordenboek.",
                ],
                "reflectie": [
                    "Welke ‘alledaagse’ term bleek moreel beladener dan gedacht?",
                    "Welke term zou je liever níet in beleid zien — en waarom?",
                ],
            },
        ],
    },
    "en": {
        "filename": "MV_21_MoreelWoordenboek_Werkvormen_EN.docx",
        "title": "Work forms for the Moral Dictionary",
        "subtitle": "MV_21  ·  Step Seeing  ·  Community of Moral Craftsmanship",
        "intro": (
            "The Moral Dictionary collects concepts that often appear in policy, codes of conduct, "
            "and conversations about integrity and social safety. These four work forms help make that "
            "language visible, discussable and testable — including through the lens of espoused theory "
            "versus theory-in-use (Argyris & Schön)."
        ),
        "forms": [
            {
                "title": "1. Word hunt in policy",
                "meta": "20–30 min  ·  individual or pairs  ·  Seeing",
                "doel": "Make visible which moral words a policy or conduct text uses most.",
                "stappen": [
                    "Choose a short passage from a code of conduct, integrity policy or team agreement (½–1 page).",
                    "Highlight all words that appear in the Moral Dictionary — or that feel morally charged even if they are not listed yet.",
                    "Note the top 5 most frequent or most loaded terms.",
                    "Discuss: which moral frame does this text speak? Rules, care, courage, safety, inclusion…?",
                ],
                "reflectie": [
                    "Which words do you miss that do belong in your practice?",
                    "Which term would you remove to make the text more honest?",
                ],
            },
            {
                "title": "2. Espoused vs. in use",
                "meta": "30–40 min  ·  3–6 people  ·  Seeing → Feeling",
                "doel": "Explore the gap between what a text claims (espoused) and what practice shows (theory-in-use).",
                "stappen": [
                    "Everyone picks 3 dictionary terms that ‘on paper’ belong to their organisation or programme.",
                    "For each term, write one concrete example where the value is visible, and one where it is missing or strained.",
                    "In the group, share only the strained examples — without identifiable details.",
                    "Together, write one sentence: ‘We say X, but in practice we often reward Y.’",
                ],
                "reflectie": [
                    "Which term turned out to be espoused but rarely in use?",
                    "What one small behavioural change would shrink the gap?",
                ],
            },
            {
                "title": "3. Cluster debate",
                "meta": "25–35 min  ·  4–8 people  ·  Weighing",
                "doel": "Discover which moral cluster pulls hardest in a situation or text — and what falls out of view.",
                "stappen": [
                    "The facilitator names a familiar situation or reads a short policy passage.",
                    "The group splits across clusters: integrity/compliance, care/relationship, courage/responsibility, safety, diversity/inclusion, transparency/trust.",
                    "Each cluster argues for 2 minutes why ‘their’ frame should dominate here.",
                    "Then: which cluster was underrepresented, and what does that cost morally?",
                ],
                "reflectie": [
                    "Which cluster usually wins in your documents?",
                    "Which cluster wins in your daily decisions?",
                ],
            },
            {
                "title": "4. Lexicon expansion",
                "meta": "15–20 min  ·  plenary or small groups  ·  Seeing",
                "doel": "Keep the dictionary alive: propose missing terms that matter in policy or practice.",
                "stappen": [
                    "Everyone notes 1–2 moral words that matter in practice but are not (yet) in the lexicon.",
                    "Per term: a one-sentence working definition + why the term is needed.",
                    "The group votes: include, discuss later, or reject.",
                    "Approved terms go to the Moral Dictionary editorial process.",
                ],
                "reflectie": [
                    "Which everyday term turned out more morally charged than expected?",
                    "Which term would you rather not see in policy — and why?",
                ],
            },
        ],
    },
}


def set_run(run, *, size=11, bold=False, italic=False, color=C_BODY):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_para(doc, text, *, size=11, bold=False, italic=False, color=C_BODY, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def build(lang: str) -> Path:
    data = CONTENT[lang]
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(data["title"])
    set_run(run, size=22, bold=True, color=C_TITLE)

    add_para(doc, data["subtitle"], size=10, color=C_ACCENT, space_after=12)
    add_para(doc, data["intro"], size=11, space_after=16)

    for form in data["forms"]:
        add_para(doc, form["title"], size=14, bold=True, color=C_TITLE, space_after=4)
        add_para(doc, form["meta"], size=10, italic=True, color=C_ACCENT, space_after=6)
        add_para(doc, f"Doel / Aim: {form['doel']}", size=11, space_after=6)
        add_para(doc, "Aanpak / Approach", size=11, bold=True, color=C_TITLE, space_after=4)
        for i, step in enumerate(form["stappen"], 1):
            add_para(doc, f"{i}. {step}", size=11, space_after=3)
        add_para(doc, "Reflectie / Reflection", size=11, bold=True, color=C_TITLE, space_after=4)
        for item in form["reflectie"]:
            add_para(doc, f"• {item}", size=11, space_after=3)
        add_para(doc, "", size=8, space_after=10)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / data["filename"]
    doc.save(out)
    return out


def main() -> None:
    for lang in ("nl", "en"):
        path = build(lang)
        print(f"Wrote {path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
